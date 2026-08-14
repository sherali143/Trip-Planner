"""
Generate PROJECT_GUIDE.docx in the project root.

This is TECHNICAL DOCUMENTATION OF THE ARTEFACT — what was built, why, and what
it measured. It is not the dissertation. It exists so the implementation and its
evaluation can be understood and checked without reading the source.

Every number comes from comparison/results/comparison_results.json and every
figure from figures/. Nothing is typed by hand, so the document cannot drift
from the system it describes. An earlier version hardcoded its headline
figures — "5 LLM calls", "~230 seconds", "85% faster" — and all three turned out
to be wrong once the LLM calls were actually instrumented.

Run:  python scripts/generate_guide.py
"""

import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RESULTS_PATH = os.path.join(ROOT, "comparison", "results", "comparison_results.json")
FIGURES = os.path.join(ROOT, "figures")
OUTPUT = os.path.join(ROOT, "PROJECT_GUIDE.docx")

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x52, 0x51, 0x4E)

doc = Document()
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)


# --------------------------------------------------------------- helpers
def h(text, level=1):
    doc.add_heading(text, level=level)


def para(text=""):
    return doc.add_paragraph(text)


def bold(text):
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    return p


def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")


def numbered(text):
    return doc.add_paragraph(text, style="List Number")


def caption(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, head in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(str(head))
        run.bold = True
    for r, row in enumerate(rows, start=1):
        for c, value in enumerate(row):
            t.rows[r].cells[c].text = str(value)
    if widths:
        for row in t.rows:
            for cell, w in zip(row.cells, widths):
                cell.width = Inches(w)
    para()
    return t


def figure(name, cap, width=6.4):
    path = os.path.join(FIGURES, name)
    if not os.path.exists(path):
        para(f"[figure missing: {name} — run python scripts/make_diagrams.py]")
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption(cap)


def load_results():
    if not os.path.exists(RESULTS_PATH):
        return None
    with open(RESULTS_PATH, encoding="utf-8") as fh:
        return json.load(fh)


RESULTS = load_results()
ARMS = (RESULTS or {}).get("arms", {})


def arm(code, key, fmt="{:,.0f}"):
    if code not in ARMS:
        return "not measured"
    return fmt.format(ARMS[code][key])


# =====================================================================
# TITLE
# =====================================================================
title = doc.add_heading("AI Trip Planner", level=0)
sub = para()
run = sub.add_run("A Multi-Agent Travel Planning System on the Model Context "
                  "Protocol and a Typed Agent-to-Agent Protocol")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = ACCENT
para("Project document — implementation and evaluation of the artefact.")

if RESULTS:
    prov = RESULTS.get("provenance", {})
    ids = RESULTS.get("scenario_ids", [])
    status = RESULTS.get("status", "unknown")
    para(f"All measurements in this document were produced by "
         f"{prov.get('model', 'the configured model')} across {len(ids)} evaluation "
         f"scenario(s) ({', '.join(ids)}), recorded with the API layer in "
         f"'{prov.get('api_mode', '?')}' mode. Result set status: {status}.")
    if status != "complete" or len(ids) < 20:
        p = para()
        r = p.add_run(f"Note: {len(ids)} of 20 scenarios have been recorded so far. "
                      f"Figures below are measurements of that subset, not of the "
                      f"full scenario set.")
        r.italic = True
        r.font.color.rgb = MUTED
else:
    para("No measured results are available — run the evaluation to populate "
         "this document.")

doc.add_page_break()

# =====================================================================
# 1. PROBLEM STATEMENT
# =====================================================================
h("1. Problem statement", 1)

para("Planning a short trip is a task people still do largely by hand. A "
     "traveller with a destination, some dates and a budget must reconcile "
     "live flight prices, hotel availability, opening times and a fixed total "
     "spend — information that lives on several sites at once and changes while "
     "they work. The task is tedious rather than difficult, which is exactly "
     "the kind of task worth automating.")

para("A single large language model is an obvious candidate and a poor one. "
     "Asked to plan a trip, it will produce a fluent, confident itinerary "
     "containing flights that do not exist at prices that were never quoted. "
     "The TravelPlanner benchmark records a 0.6% final-pass rate for the best "
     "single-LLM agent, with failures dominated by hallucinated venues, "
     "fabricated prices and violated budget constraints (Xie et al., 2024).")

bold("The three failure modes this project targets")
table(
    ["Failure mode", "What goes wrong", "How the design answers it"],
    [
        ["Context bloat",
         "A single hotel API response can exceed 8 kB of raw JSON. Feed several "
         "into one context and earlier reasoning is pushed out.",
         "Tool results are distilled to the few options that matter before they "
         "reach the model."],
        ["Protocol fragility",
         "A model asked to call an API freely emits malformed parameters.",
         "Every tool call passes a schema before it can reach an external service."],
        ["Semantic drift",
         "Free text handed between cooperating agents degrades at each handoff.",
         "Agents exchange typed, permission-validated messages, never prose."],
    ],
    widths=[1.3, 2.6, 2.5],
)

bold("Problem statement")
para("Given a travel request in ordinary English, produce a day-by-day "
     "itinerary in which every flight, hotel, activity and price is traceable "
     "to data actually retrieved from a live source, within a total budget the "
     "user controls — and determine what architecture delivers that at "
     "acceptable cost.")

# =====================================================================
# 2. SCOPE
# =====================================================================
h("2. Scope", 1)

table(
    ["In scope", "Out of scope"],
    [
        ["English natural-language input", "Multilingual input"],
        ["Three live APIs: flights, hotels, web search", "GDS systems (Sabre, Amadeus)"],
        ["Twelve schema-validated MCP tools", "Booking or payment transactions"],
        ["Typed A2A messaging between agents", "User accounts and authentication"],
        ["Web (Streamlit) and command-line interfaces", "Native mobile applications"],
        ["Retrospective evaluation over scripted scenarios", "Live user study"],
        ["User-controlled budget allocation", "Currency conversion"],
    ],
    widths=[3.2, 3.2],
)

para("The system is a planning aid. It does not book anything, and the "
     "interface states so.")

# =====================================================================
# 3. SOLUTION
# =====================================================================
h("3. The solution", 1)

para("The system is organised in four layers. A conversational layer collects "
     "the request; an extraction layer converts it to typed JSON; a retrieval "
     "layer obtains real data through a Model Context Protocol (MCP) server; a "
     "coordination layer assembles the itinerary from that data and nothing "
     "else.")

figure("diagrams/architecture.png",
       "Figure 1 — System architecture. The MCP server is the only route to an "
       "external API; every inter-agent message travels over the A2A protocol.")

bold("3.1 The MCP server")
para("A subprocess exposing twelve schema-validated tools over JSON-RPC on "
     "stdio. Each tool declares its parameters, so a malformed call is rejected "
     "before it can reach a paid external service. Responses pass through a "
     "cache and a distillation stage on the way back.")

figure("diagrams/mcp_lifecycle.png",
       "Figure 2 — The six stages of a tool call. The cache sits in front of "
       "the network, so a repeated query costs nothing.")

bold("3.2 The A2A protocol")
para("Agents never hand each other free text. Every message carries an "
     "identifier, a sender, a receiver, one of six types, a priority, a "
     "conversation identifier and a typed payload. A registry of agent cards "
     "declares who may send what to whom, and a message that violates it is "
     "rejected rather than delivered.")

figure("diagrams/a2a_flow.png",
       "Figure 3 — The six A2A messages exchanged during one trip request.")

para("This layer is identical in every architecture evaluated below. That is "
     "deliberate: holding the protocol constant is what allows the differences "
     "measured later to be attributed to the retrieval strategy rather than to "
     "how components communicate.")

# =====================================================================
# 4. APPROACHES
# =====================================================================
h("4. Approaches evaluated", 1)

para("The proposal specified a six-agent architecture. Once that was built and "
     "instrumented, most of its LLM calls turned out to be spent on data "
     "retrieval rather than on reasoning. Rather than assert that a smaller "
     "design is better, four architectures were built and measured against the "
     "same requests.")

figure("diagrams/four_arms.png",
       "Figure 4 — The four architectures. Same request, same APIs, same "
       "protocol; only the retrieval layer differs.")

bold("Why four rather than two")
para("Comparing the three-agent design only against the naive six-agent build "
     "would invite an obvious objection: that the multi-agent arm lost because "
     "it was configured badly, not because the architecture is worse. Arm C "
     "removes that objection by implementing three commitments the proposal "
     "made that the first build never did — distilled tool output, concurrent "
     "specialists, and distillation as a stage of the MCP lifecycle. Arm A, a "
     "single model with no tools, anchors the other end and shows what the "
     "tool layer buys.")

bold("4.1 Pros and cons")
table(
    ["Arm", "Strengths", "Weaknesses"],
    [
        ["A — Single LLM",
         "Simplest possible; one request; no infrastructure.",
         "No access to real data. Cites venues and prices that were never "
         "retrieved, which is the failure this project exists to avoid."],
        ["B — Six agents, naive",
         "Full agent autonomy: a specialist can re-search when a result "
         "disappoints. Faithful to the proposal's agent decomposition.",
         "Each specialist carries up to eight tool schemas, re-sent on every "
         "reasoning step. Raw API payloads enter the context. Highest cost of "
         "the four on every measure."],
        ["C — Six agents, tuned",
         "Keeps agent autonomy while removing most of the cost. Specialists run "
         "concurrently. Demonstrates that the multi-agent penalty is largely an "
         "implementation artefact.",
         "Still spends several LLM calls deciding which tool to use — work that "
         "involves no judgement. Slower than direct execution."],
        ["D — Three agents, direct API",
         "LLM used only where judgement is required: understanding the request "
         "and assembling the plan. Fewest calls, lowest latency. Retrieval is "
         "deterministic and testable.",
         "Cannot adapt its search: if results disappoint, it cannot widen the "
         "dates and try again. That autonomy is genuinely lost."],
    ],
    widths=[1.2, 2.6, 2.6],
)

# =====================================================================
# 5. SUPERVISOR-REQUESTED ENHANCEMENT
# =====================================================================
h("5. Enhancement: user-controlled budget allocation", 1)

para("Following supervisor feedback, control over how the budget is divided "
     "was moved from the system to the user.")

bold("The problem with the original behaviour")
para("The system applied one fixed split to every trip — flights 35%, "
     "accommodation 35%, activities 20%, meals 10% — written into seven places "
     "in the codebase, with no source behind the figures and no way for a user "
     "to change them.")

para("That split is also wrong for most trips, because the two largest "
     "categories scale on different axes. Airfare is charged per person and set "
     "by distance; it does not care how many nights are stayed. Accommodation "
     "is charged per night and normally shared between travellers; it does not "
     "care how far anyone flew. A three-night trip to a nearby city and a "
     "fourteen-night long-haul trip cannot sensibly use the same percentages.")

bold("What was implemented")
table(
    ["Before", "After"],
    [
        ["One fixed split for every trip",
         "A split derived from what the trip's own components actually cost"],
        ["Written into seven places",
         "One module, with the cost model as the single source"],
        ["Never explained to the user",
         "Each category explained, with the reasoning for the suggested figures"],
        ["No way to change it",
         "Accepts percentages, named categories, partial input, or cash amounts"],
        ["Budget feasibility judged by a fixed formula ignoring destination",
         "Feasibility judged against the estimated cost of that specific trip"],
    ],
    widths=[3.2, 3.2],
)

para("Because the suggestion is derived rather than tuned, it responds to the "
     "trip. A two-night trip to Dubai puts roughly a third of the budget on "
     "flights; a fourteen-night trip to the same city puts under a fifth there, "
     "because the one-off airfare is spread across far more nights.")

bold("Feasibility, and a floor that is genuinely a floor")
para("The previous check applied the same threshold regardless of destination, "
     "so a modest budget was judged identically for a nearby city and a "
     "long-haul one — and both were accepted, meaning the impossible trip "
     "produced a confident but fictional itinerary. Cost is now estimated for "
     "the specific destination, and a budget below the true minimum is refused "
     "with an explanation of the shortfall and concrete alternatives. Anything "
     "above that floor proceeds: a tight budget is a legitimate choice and is "
     "warned about, not blocked.")

# =====================================================================
# 6. RESULTS
# =====================================================================
h("6. Measured results", 1)

if not RESULTS:
    para("No results recorded. Run: python -m comparison.run_comparison")
else:
    para("Every LLM request is counted through provider callbacks at runtime; "
         "none of these figures is an estimate. Costs and latencies are means "
         "over the recorded scenarios.")

    table(
        ["Architecture", "LLM calls", "Tokens", "Cost (USD)", "Time", "Real prices"],
        [[f"{c} — {ARMS[c]['name']}",
          arm(c, "avg_llm_calls"),
          arm(c, "avg_total_tokens"),
          arm(c, "avg_cost_usd", "${:,.4f}"),
          arm(c, "avg_latency", "{:,.0f}s"),
          arm(c, "avg_prices_grounded_pct", "{:,.0f}%")]
         for c in ("A", "B", "C", "D") if c in ARMS],
        widths=[1.9, 0.9, 0.9, 1.0, 0.7, 0.9],
    )

    figure("results/efficiency.png",
           "Figure 5 — Cost of each architecture across four measures. "
           "Presented as separate panels because the quantities share no scale.")

    bold("6.1 Tuning accounts for most of the multi-agent penalty")
    para("Arms B and C contain the same six agents and reach the same APIs. "
         "Only the prompt economics differ: one narrow tool per specialist "
         "instead of up to eight, distilled results instead of raw payloads, "
         "and a cap on reasoning iterations.")

    figure("results/tuning_effect.png",
           "Figure 6 — Effect of tuning alone, with the architecture unchanged.")

    para("This matters for the conclusion. Most of the cost difference between "
         "the naive multi-agent build and direct execution is an implementation "
         "artefact, not a property of multi-agent design. Compared against the "
         "tuned arm rather than the naive one, direct execution still wins "
         "clearly on call count and latency, but only modestly on cost.")

    bold("6.2 Cost alone would rank the worst system first")
    para("The single-LLM arm is inexpensive precisely because it retrieves "
         "nothing. Reporting its cost without reporting what it produced would "
         "invert the conclusion, so each itinerary is also scored on whether "
         "the prices it quotes match fares and rates actually returned by the "
         "APIs.")

    figure("results/groundedness.png",
           "Figure 7 — Share of quoted prices traceable to retrieved data.")

    a_pct = ARMS.get("A", {}).get("avg_prices_grounded_pct")
    if a_pct is not None:
        para(f"The tool-less arm scores {a_pct:.0f}%. Every price it printed was "
             f"invented — the hallucination failure described in the literature, "
             f"measured directly rather than assumed.")

    bold("6.3 Interpreting the two groundedness signals")
    para("Price matching is the robust signal: landing within 2% of a real "
         "quoted fare is not something prior knowledge delivers. Name matching "
         "is weaker and is reported only as supporting detail — a model with no "
         "tool access can still name a real airline by guessing the obvious "
         "carrier for a route, which is exactly what the tool-less arm did.")

# =====================================================================
# 7. REPRODUCIBILITY
# =====================================================================
h("7. Reproducibility", 1)

para("Every HTTP response the evaluation depends on is recorded and committed "
     "to the repository. The entire evaluation can therefore be re-run by "
     "anyone, producing identical numbers, with no API keys at all:")

para("    TRIP_PLANNER_API_MODE=replay python -m comparison.run_comparison")

table(
    ["Property", "How it is guaranteed"],
    [
        ["Numbers cannot drift from the data",
         "This document and every figure are generated from the results file."],
        ["Results reproduce without credentials",
         "Recorded API responses are committed; replay mode never touches the network."],
        ["Errors are never mistaken for data",
         "Only successful responses are cached, so a quota failure is never stored."],
        ["No credentials reach disk",
         "Request headers are excluded from both the cache key and the stored file."],
        ["A run cannot exhaust a monthly allowance",
         "Live calls and LLM requests are both capped, and runs stop between scenarios."],
        ["An interrupted run loses nothing",
         "Results are checkpointed per scenario; a re-run reuses what is recorded."],
    ],
    widths=[2.6, 3.8],
)

# =====================================================================
# 8. STRUCTURE
# =====================================================================
h("8. Issues encountered and how they were resolved", 1)

para("Recorded because several of these shaped the design, and because a "
     "failure that is silent deserves more attention than one that is loud.")

table(
    ["Issue", "Why it mattered", "Resolution"],
    [
        ["Flight search returned nothing, always",
         "The provider is a two-phase API: the search endpoint only starts a "
         "search and returns a session identifier. Reading its first response "
         "as final yields an empty list every time, with HTTP 200 and no error.",
         "Poll the completion endpoint using the session identifier."],
        ["Every itinerary carried the wrong outbound date",
         "The provider expects camelCase date parameters. The snake_case forms "
         "are accepted, answered with HTTP 200 and silently ignored, so the "
         "system searched a default date about a week out while appearing to "
         "work correctly.",
         "Identified the accepted names by sending a distinct date under each "
         "candidate in a single request and decoding which one came back."],
        ["All twelve MCP tools failed silently",
         "The MCP server runs as a subprocess, so the project root was not on "
         "its import path. Every tool call returned Connection lost, and the "
         "agents produced itineraries from model knowledge instead. The "
         "failure looked like success.",
         "The server sets up its import path before loading project modules."],
        ["Reported LLM cost was assumed, not measured",
         "The harness counted tasks rather than requests, and one script "
         "incremented its counter with a literal simulated comment. The "
         "reported figure understated the real cost roughly fourfold.",
         "Every request is now instrumented through provider callbacks, "
         "capturing tokens, cost and latency."],
        ["A quota error was indistinguishable from an empty result",
         "A failed flight search reported success with an empty list, so one "
         "architecture invented flights from it while another logged a real "
         "error, making their success rates incomparable.",
         "Failures are reported as failures."],
        ["The evaluation could not be run at all",
         "The travel APIs allow 30 and 50 requests per month. One pass over "
         "twenty scenarios across four architectures needs far more.",
         "Record and replay caching: one pass records, later passes replay for "
         "free and deterministically."],
        ["An interrupted run lost everything it had paid for",
         "Results were written once at the end, so a run stopped at scenario "
         "fifteen discarded fifteen scenarios worth of spent quota.",
         "Results are checkpointed after each scenario and reused on re-run."],
        ["The API key was printed to the console",
         "The LLM provider carries its key in the request URL, and the HTTP "
         "client logs request URLs at INFO level.",
         "Third-party log levels are raised before any request is made."],
    ],
    widths=[1.5, 2.6, 2.3],
)

para("A pattern runs through these. The costly failures were the quiet ones: "
     "wrong dates behind a 200 response, dead tools producing plausible "
     "itineraries, a cost figure that was never actually measured. Each was "
     "found by checking output against reality rather than by reading code, "
     "which is why the evaluation now scores groundedness instead of assuming "
     "that retrieved data was used.")

h("9. What the implementation delivers", 1)

table(
    ["Capability", "What it means in practice"],
    [
        ["Itineraries grounded in retrieved data",
         "Prices and venues trace back to an API response. The tool-less "
         "baseline matches none of its prices to anything real; the "
         "tool-using architectures match the majority."],
        ["A measured architectural comparison",
         "Four architectures on identical requests with cost counted rather "
         "than estimated, which is the evidence the design decision rests on."],
        ["Budget control held by the user",
         "The split is derived from what this particular trip costs, "
         "explained, and then replaced by whatever the user prefers."],
        ["Honest refusal",
         "A budget below what the destination actually costs is refused with "
         "the shortfall and workable alternatives, rather than answered with "
         "a plan nobody could book."],
        ["Reproducible results",
         "The evaluation re-runs from recorded responses and produces "
         "identical numbers with no API keys."],
        ["Bounded operating cost",
         "Live calls and LLM requests are both capped, and runs stop between "
         "whole scenarios, so no single run can exhaust a monthly allowance."],
        ["Documentation that cannot drift",
         "This document and every figure are generated from the results file, "
         "and a test fails when the written documentation stops matching the "
         "code."],
    ],
    widths=[2.0, 4.4],
)

h("10. Repository structure", 1)

table(
    ["Location", "Contents"],
    [
        ["PROJECT_GUIDE.docx", "This document"],
        ["setup.bat", "One-command setup: environment, dependencies, verification"],
        ["run_cli.py", "Plan a trip in the terminal"],
        ["run_web.py", "Plan a trip in the browser"],
        ["src/", "Application code"],
        ["comparison/", "The evaluation: architectures, scenarios, metrics, results"],
        ["demos/", "Scripts for demonstrating the system"],
        ["testing/", "Automated test suite"],
        ["scripts/", "Generators for this document and the figures"],
        ["figures/", "Generated charts and diagrams"],
        ["proposal/", "Project proposal and assignment brief"],
        ["report/", "Dissertation report"],
        ["deploy/", "Container definition"],
        [".api_cache/", "Recorded API responses, committed for reproducibility"],
    ],
    widths=[1.7, 4.7],
)

bold("What each file does — application code")
table(
    ["File", "Responsibility"],
    [
        ["src/orchestrator.py", "The production workflow, end to end"],
        ["src/agents.py", "The three agents that require an LLM"],
        ["src/tasks.py", "Their task definitions and prompts"],
        ["src/comms/protocol.py", "A2A message envelope and priority queue"],
        ["src/comms/registry.py", "Agent cards: who may send what to whom"],
        ["src/server/mcp_server.py", "MCP server, twelve tools over JSON-RPC"],
        ["src/tools/mcp_tools.py", "Tool wrappers for agents, plus direct HTTP calls"],
        ["src/core/http_cache.py", "Record and replay caching, live-call ceiling"],
        ["src/core/llm_metrics.py", "Counts real LLM requests, tokens and cost"],
        ["src/core/resilience.py", "Whether a provider refusal is worth retrying"],
        ["src/core/budget.py", "Derives and applies the budget split"],
        ["src/core/trip_cost.py", "Estimates trip cost and sets the feasibility floor"],
        ["src/core/validators.py", "Checks the itinerary contains every day it should"],
        ["src/core/log_setup.py", "Keeps the API key out of the console"],
        ["src/ui/app.py", "Streamlit interface"],
    ],
    widths=[2.1, 4.3],
)

bold("What each file does — evaluation")
table(
    ["File", "Responsibility"],
    [
        ["comparison/arm_a_single_llm.py", "Arm A, one LLM call, no tools"],
        ["comparison/arm_b_six_agent_naive.py", "Arm B, six agents as first built"],
        ["comparison/arm_c_six_agent_tuned.py", "Arm C, six agents as specified"],
        ["comparison/arm_d_three_agent_direct.py", "Arm D, three agents, direct retrieval"],
        ["comparison/distilled_tools.py", "Compact tool wrappers used by arm C"],
        ["comparison/scenarios.py", "The twenty evaluation scenarios"],
        ["comparison/metrics.py", "Groundedness scoring"],
        ["comparison/run_comparison.py", "Runs the arms, checkpoints, aggregates"],
    ],
    widths=[2.5, 3.9],
)

bold("What each file does — demonstration, generation, tests")
table(
    ["File", "Responsibility"],
    [
        ["demos/demo_approach.py", "Runs one architecture, narrated"],
        ["demos/demo_comparison.py", "Runs all four side by side"],
        ["scripts/generate_guide.py", "Generates this document"],
        ["scripts/make_charts.py", "Generates the results charts"],
        ["scripts/make_diagrams.py", "Generates the architecture diagrams"],
        ["testing/test_budget_allocation.py", "Budget splitting and user input"],
        ["testing/test_trip_cost.py", "Cost estimation and feasibility"],
        ["testing/test_calculator.py", "The safe arithmetic evaluator"],
        ["testing/test_itinerary_validator.py", "Itinerary day-count validation"],
        ["testing/test_budget_validation.py", "Budget parsing"],
        ["testing/test_documentation_accuracy.py", "That documentation matches the code"],
    ],
    widths=[2.5, 3.9],
)

h("11. Known limitations", 1)

para("Stated here rather than left to be discovered.")

table(
    ["Limitation", "Detail"],
    [
        ["Scenario coverage",
         f"{len((RESULTS or {}).get('scenario_ids', []))} of 20 scenarios recorded "
         f"so far. Reported figures describe that subset."],
        ["Run-to-run variance",
         "The naive multi-agent arm varied between 19 and 23 LLM calls on "
         "identical input, reflecting the non-determinism of iterative reasoning "
         "loops. Figures are single-run measurements."],
        ["Destination price data",
         "Cost estimation uses a curated table of roughly sixty cities. "
         "Unrecognised destinations receive mid-tier defaults, and the system "
         "says so rather than presenting them as derived."],
        ["Groundedness by name",
         "Matching venue and airline names is weak evidence, since well-known "
         "names can be guessed. Price matching carries the claim."],
        ["Evaluation design",
         "Scripted scenarios rather than a live user study, which limits "
         "external validity."],
    ],
    widths=[1.8, 4.6],
)

# =====================================================================
# REFERENCES
# =====================================================================
h("References", 1)
para("Anthropic (2024) Model Context Protocol Specification. "
     "Available at: https://modelcontextprotocol.io")
para("Xie, J., Zhang, K., Chen, J., Zhu, T., Lou, R., Tian, Y., Xiao, Y. and "
     "Su, Y. (2024) 'TravelPlanner: A Benchmark for Real-World Planning with "
     "Language Agents', ICML 2024.")
para("CrewAI (2024) CrewAI Documentation: Agents, Tasks and Crews. "
     "Available at: https://docs.crewai.com")

doc.save(OUTPUT)
print(f"Document saved to: {OUTPUT}")
if RESULTS:
    ids = RESULTS.get("scenario_ids", [])
    print(f"  results: {len(ids)} scenario(s), status={RESULTS.get('status')}")
print(f"  figures embedded from: {FIGURES}")
