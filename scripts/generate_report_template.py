"""
Generate report/Dissertation_TEMPLATE.docx — a formatted shell to write into.

This produces STRUCTURE, not content. Every section contains a guidance box
saying what belongs there and which artefact in this repository supports it,
followed by an empty paragraph to type into. The guidance boxes are styled in
grey italics and marked for deletion, so they cannot be mistaken for submission
text.

No analytical prose is generated. The assignment brief states that assessed
writing must be the student's own and not produced by an AI tool, so the
argument, the analysis and the conclusions are left blank deliberately.

Formatting follows the brief: 11 pt body text, 1.5 line spacing, student number
only on the cover (the submission is marked anonymously).

Run:  python scripts/generate_report_template.py
"""

import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt, RGBColor

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
FIGURES = os.path.join(ROOT, "figures")
RESULTS_PATH = os.path.join(ROOT, "comparison", "results", "comparison_results.json")
OUTPUT = os.path.join(ROOT, "report", "Dissertation_TEMPLATE.docx")

GUIDANCE = RGBColor(0x80, 0x80, 0x80)
FILL = RGBColor(0x1F, 0x4E, 0x79)

doc = Document()

# Brief: font size 11, 1.5 spacing, to leave room for marker annotations.
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after = Pt(6)


def guidance(lines):
    """A grey italic block telling the writer what belongs here. Delete when done."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("GUIDANCE — delete this box before submitting")
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GUIDANCE
    for line in lines:
        q = doc.add_paragraph()
        q.paragraph_format.line_spacing = 1.0
        q.paragraph_format.left_indent = Inches(0.25)
        r = q.add_run(line)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GUIDANCE
    doc.add_paragraph()


def blank(n=1):
    for _ in range(n):
        doc.add_paragraph()


def heading(text, level=1):
    return doc.add_heading(text, level=level)


def placeholder(label="Write this section here."):
    p = doc.add_paragraph()
    run = p.add_run(f"[{label}]")
    run.font.color.rgb = GUIDANCE
    run.italic = True
    return p


def figure(rel_path, caption_text, width=6.0):
    path = os.path.join(FIGURES, rel_path)
    if not os.path.exists(path):
        placeholder(f"FIGURE MISSING: figures/{rel_path} — run the generators")
        return
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(caption_text)
    run.italic = True
    run.font.size = Pt(9)
    blank()


def page_break():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def load_results():
    if not os.path.exists(RESULTS_PATH):
        return None
    with open(RESULTS_PATH, encoding="utf-8") as fh:
        return json.load(fh)


RESULTS = load_results()

# ======================================================================
# COVER
# ======================================================================
for _ in range(4):
    blank()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("[DISSERTATION TITLE]")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = FILL

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = s.add_run("[Subtitle, if used]")
r.font.size = Pt(13)
r.font.color.rgb = GUIDANCE

blank(3)

for line in [
    "[STUDENT NUMBER]",
    "",
    "CMP7200 Individual Master's Project",
    "Faculty of Computing, Engineering and the Built Environment",
    "Birmingham City University",
    "",
    "[SUBMISSION DATE]",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    if line.startswith("["):
        run.font.color.rgb = GUIDANCE
        run.bold = True

blank(2)
guidance([
    "The brief states this assessment is marked anonymously: student number only, no name.",
    "Attach the BCU cover sheet in front of this page.",
    "Body text is already set to 11 pt with 1.5 line spacing, as the brief requires.",
])

page_break()

# ======================================================================
# FRONT MATTER
# ======================================================================
heading("Abstract", 1)
guidance([
    "Not counted in the word limit. Around 250-300 words.",
    "Cover: the problem, what was built, how it was evaluated, the headline result,",
    "and what it means. Write this last, once the results are final.",
])
placeholder("Abstract")
blank()

heading("Acknowledgements", 1)
placeholder("Acknowledgements")
blank()

heading("Table of Contents", 1)
guidance([
    "In Word: References > Table of Contents > Automatic Table.",
    "It populates from the Heading styles already applied throughout this file.",
    "Update it (right-click > Update Field) after every edit.",
])
placeholder("Insert automatic table of contents")
blank()

heading("List of Figures", 1)
guidance([
    "In Word: References > Insert Table of Figures, using the caption style.",
])
placeholder("Insert table of figures")

page_break()

# ======================================================================
# CHAPTERS
# ======================================================================
CHAPTERS = [
    ("1. Introduction", "approximately 1,200 words", [
        ("1.1 Background and motivation", [
            "Why trip planning is worth automating. Keep it grounded — the reader",
            "wants the problem, not a tourism-industry essay.",
        ]),
        ("1.2 Problem statement", [
            "State it as one testable proposition.",
            "Source: PROJECT_GUIDE.docx section 1 states the problem as implemented.",
        ]),
        ("1.3 Aim and objectives", [
            "Restate from the proposal. Mark any objective that changed and say why —",
            "a marker comparing the two documents will notice, and an explained",
            "deviation reads as control while an unexplained one reads as drift.",
        ]),
        ("1.4 Scope", [
            "In scope and out of scope. Source: PROJECT_GUIDE.docx section 2.",
        ]),
        ("1.5 Contributions", [
            "What this work adds. Be specific and modest — three concrete claims",
            "beat one grand one.",
        ]),
        ("1.6 Structure of the dissertation", [
            "One short paragraph per chapter.",
        ]),
    ]),
    ("2. Literature Review", "approximately 2,500 words — LO2, 25% of the mark", [
        ("2.1 Single-LLM tool use and its limits", [
            "Toolformer, ReAct, Reflexion, and the evidence on where single-model",
            "tool use breaks down.",
        ]),
        ("2.2 Multi-agent frameworks", [
            "AutoGen, MetaGPT, CrewAI, GPTSwarm. Justify the framework you chose.",
        ]),
        ("2.3 Tool and messaging protocols", [
            "MCP, schema-validated function calling, FIPA-ACL as the antecedent for",
            "typed inter-agent messaging.",
        ]),
        ("2.4 AI travel planning systems", [
            "TravelPlanner and its 0.6% final-pass rate; prior itinerary systems;",
            "what none of them do together.",
        ]),
        ("2.5 Conceptual model and research gap", [
            "This subsection is where the marks are. The criteria distinguish",
            "'compares literature' from 'constructs a conceptual model derived from",
            "these sources' — aim for the second.",
            "For each of the three failure modes you designed against — context bloat,",
            "protocol fragility, semantic drift — name the work that identifies it and",
            "the design decision that answers it. Your measurements let you close that",
            "loop with evidence, which most projects cannot do.",
        ]),
    ]),
    ("3. Research Methodology", "approximately 2,000 words — LO3, 35% of the mark", [
        ("3.1 Research paradigm", [
            "Design Science Research (Hevner et al., 2004): build an artefact,",
            "evaluate it, learn from it. Say why it fits and what you rejected.",
        ]),
        ("3.2 Why four architectures", [
            "The most important argument in the chapter.",
            "Source: comparison/README.md justifies each arm individually.",
            "Land this point: arm C exists so the comparison is not against a straw",
            "man. A marker could otherwise object that the multi-agent arm lost",
            "through misconfiguration. Making that argument yourself is worth more",
            "than the result it protects.",
            "State plainly that the six-agent arms instantiate five agents — the",
            "conversational agent is omitted so every arm receives an identical",
            "request — and call it a five-agent ablation of the six-agent design.",
        ]),
        ("3.3 Scenario design", [
            "Twenty scenarios spanning trip length, distance, party size, destination",
            "price tier and budget, including deliberately infeasible ones.",
            "Source: the docstring in comparison/scenarios.py explains the axes.",
        ]),
        ("3.4 Metrics", [
            "Cost, latency, and groundedness. Explain why cost alone would rank the",
            "worst system first, and why price matching is stronger evidence than",
            "name matching.",
            "Source: comparison/metrics.py documents both, including the limitation.",
        ]),
        ("3.5 Controls and fairness", [
            "Identical request strings, one shared retry policy, LLM usage counted",
            "rather than estimated, budget allocation held constant across arms.",
            "Source: comparison/run_comparison.py and its README.",
        ]),
        ("3.6 Reproducibility", [
            "Record and replay: the evaluation re-runs from committed responses with",
            "no API keys. Source: PROJECT_GUIDE.docx section 7.",
        ]),
        ("3.7 Ethics and limitations of the method", [
            "Read-only API use within free tiers, no personal data, planning aid",
            "rather than a booking system, scripted scenarios rather than a user study.",
        ]),
    ]),
    ("4. Design and Implementation", "approximately 2,800 words — LO4, 40% of the mark", [
        ("4.1 System architecture", [
            "Walk through the four layers. The figure is inserted below — reference it",
            "in your text as Figure 4.1 and explain what the reader should take from it.",
        ]),
        ("4.2 The MCP server", [
            "Twelve schema-validated tools over JSON-RPC on stdio, and the six-stage",
            "lifecycle of a call.",
        ]),
        ("4.3 The A2A protocol", [
            "Envelope, six message types, agent cards, permission validation.",
            "Make the point that this layer is identical in every arm — that is what",
            "lets the comparison attribute differences to retrieval strategy.",
        ]),
        ("4.4 The four architectures", [
            "How each one differs, and what each gives up.",
        ]),
        ("4.5 Budget allocation and feasibility", [
            "Your supervisor's requested change.",
            "Cover why one fixed split was wrong: airfare scales per person and by",
            "distance, accommodation scales per night and is shared, so a short",
            "long-haul trip and a long regional trip cannot use the same percentages.",
            "Then what replaced it, and that the feasibility floor is now",
            "destination-specific rather than one hardcoded threshold.",
            "Source: PROJECT_GUIDE.docx section 5.",
        ]),
        ("4.6 Implementation problems and their resolution", [
            "Give this real space — it is among the strongest material available.",
            "The valuable cases are the quiet failures: dates silently ignored behind",
            "an HTTP 200, twelve tools returning 'Connection lost' while the system",
            "produced confident itineraries anyway, a reported cost that had never",
            "been measured. Each shows diagnosis rather than coding.",
            "Source: PROJECT_GUIDE.docx section 8 lists them with resolutions.",
        ]),
        ("4.7 Testing", [
            "What is covered, what is deliberately not (nothing calls a live API),",
            "and why the documentation itself is tested.",
        ]),
    ]),
    ("5. Evaluation and Results", "approximately 2,200 words", [
        ("5.1 Experimental setup", [
            "Model, API mode, how many scenarios were recorded, and the date.",
            "Check 'status' and 'scenario_ids' in comparison/results/",
            "comparison_results.json and state the coverage honestly here. If the run",
            "is partial, say so — do not let a reader assume twenty.",
        ]),
        ("5.2 Cost across architectures", [
            "Interpret the figure below. Do not just restate the numbers — say what",
            "they mean and why the pattern occurs.",
        ]),
        ("5.3 The effect of tuning", [
            "Arms B and C contain the same six agents and reach the same APIs; only",
            "the prompt economics differ. This is the most interesting result you have.",
        ]),
        ("5.4 Groundedness", [
            "The counterweight to cost. Explain what a zero here means.",
        ]),
        ("5.5 Threats to validity", [
            "Cover all five, in your own words:",
            "  run-to-run variance (arm B varied 19-23 calls on identical input,",
            "  so the figures are single-run measurements);",
            "  scenario coverage;",
            "  destination price data (a ~60-city table; unknown cities get defaults);",
            "  name-based groundedness being weak evidence;",
            "  scripted scenarios rather than a live user study.",
            "Naming these yourself is worth more than hoping they go unnoticed.",
        ]),
    ]),
    ("6. Discussion", "approximately 900 words", [
        ("6.1 Interpretation of the results", [
            "Argue the narrower claim, because it is the one the data supports:",
            "tuning removed most of the multi-agent penalty, so the surviving finding",
            "is about call count and latency rather than a large cost difference.",
            "That is harder to attack than 'fewer agents win'.",
        ]),
        ("6.2 Relation to the literature", [
            "Tie the groundedness result back to the hallucination findings in",
            "chapter 2. This is where LO2 marks are earned a second time.",
        ]),
        ("6.3 Deviations from the proposal", [
            "Kiwi.com to fly-scraper, GPT-4o to Gemini, thirteen tools to twelve,",
            "six agents to three in production. Each with its reason.",
            "Source: proposal/README.md tabulates them.",
        ]),
    ]),
    ("7. Conclusion and Future Work", "approximately 400 words", [
        ("7.1 Objectives revisited", [
            "Take each objective from chapter 1 and say whether it was met.",
        ]),
        ("7.2 Contributions", [
            "Restate briefly. No new material here.",
        ]),
        ("7.3 Future work", [
            "Ground each in a real limitation: deriving destination price tiers from",
            "live data rather than a fixed table; a hybrid architecture that uses an",
            "agent only when a first search fails; a live user study.",
        ]),
    ]),
]

FIGURE_SLOTS = {
    "4.1 System architecture": ("diagrams/architecture.png",
                                "Figure 4.1 — System architecture."),
    "4.2 The MCP server": ("diagrams/mcp_lifecycle.png",
                           "Figure 4.2 — MCP tool call lifecycle."),
    "4.3 The A2A protocol": ("diagrams/a2a_flow.png",
                             "Figure 4.3 — A2A message flow for one trip."),
    "4.4 The four architectures": ("diagrams/four_arms.png",
                                   "Figure 4.4 — The four architectures compared."),
    "5.2 Cost across architectures": ("results/efficiency.png",
                                      "Figure 5.1 — Cost of each architecture."),
    "5.3 The effect of tuning": ("results/tuning_effect.png",
                                 "Figure 5.2 — Effect of tuning (arm B to arm C)."),
    "5.4 Groundedness": ("results/groundedness.png",
                         "Figure 5.3 — Prices traceable to retrieved data."),
}

for chapter_title, budget, sections in CHAPTERS:
    heading(chapter_title, 1)
    p = doc.add_paragraph()
    run = p.add_run(budget)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GUIDANCE
    blank()

    for section_title, notes in sections:
        heading(section_title, 2)
        guidance(notes)
        placeholder()
        if section_title in FIGURE_SLOTS:
            rel, cap = FIGURE_SLOTS[section_title]
            figure(rel, cap)
        blank()

    page_break()

# ======================================================================
# BACK MATTER
# ======================================================================
heading("References", 1)
guidance([
    "BCU Harvard. Not counted in the word limit.",
    "Carry over everything cited in the proposal, plus whatever chapter 2 adds.",
])
placeholder("References")
page_break()

heading("Appendices", 1)
guidance([
    "Not counted in the word limit. Candidates:",
    "  A — The twenty evaluation scenarios (comparison/scenarios.py)",
    "  B — Full results table (comparison/results/comparison_results.json)",
    "  C — Repository structure and file listing (PROJECT_GUIDE.docx section 10)",
    "  D — Link to the repository",
])
placeholder("Appendices")

os.makedirs(os.path.join(ROOT, "report"), exist_ok=True)
doc.save(OUTPUT)

print(f"Template saved to: {OUTPUT}")
print(f"  chapters: {len(CHAPTERS)}")
print(f"  sections: {sum(len(s) for _, _, s in CHAPTERS)}")
print(f"  figures placed: {len(FIGURE_SLOTS)}")
if RESULTS:
    ids = RESULTS.get("scenario_ids", [])
    print(f"  results available: {len(ids)} scenario(s), status={RESULTS.get('status')}")
print("  formatting: 11 pt, 1.5 line spacing, student number only on cover")
