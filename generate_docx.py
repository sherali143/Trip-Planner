"""Generate dissertation explanation DOCX"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# ============================================
# STYLES
# ============================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def add_para(text):
    return doc.add_paragraph(text)

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r+1].cells[c].text = str(val)
    return table

# ============================================
# PAGE 1: PROJECT OVERVIEW
# ============================================
add_heading('AI Trip Planner — Project Overview', level=0)

add_para('')
add_bold_para('What is this project?')
add_para(
    'This is an AI-powered trip planner that creates complete travel itineraries. '
    'You tell it where you want to go, when, and your budget — and it produces '
    'a day-by-day plan with flights, hotels, attractions, and restaurants.'
)

add_bold_para('Scope')
add_bullet('Takes user input in plain English (e.g., "Plan a trip to Istanbul for 4 nights")')
add_bullet('Fetches real-time data from live APIs — flights (fly-scraper), hotels (Booking.com), attractions and restaurants (Serper web search)')
add_bullet('Uses AI agents (powered by Google Gemini) to understand requests and assemble itineraries')
add_bullet('Runs in two modes: Terminal (CLI) and Web UI (Streamlit)')
add_bullet('Built as a dissertation project comparing two architectures')

add_bold_para('Technology Stack')
add_table(
    ['Component', 'Technology'],
    [
        ['AI Agents', 'CrewAI + Google Gemini 2.5 Flash (via LiteLLM)'],
        ['Agent Communication', 'A2A Protocol (custom implementation)'],
        ['Tool Integration', 'MCP Server (Model Context Protocol)'],
        ['Flight Data', 'fly-scraper API (via RapidAPI)'],
        ['Hotel Data', 'Booking.com API (via RapidAPI)'],
        ['Web Search', 'Serper API'],
        ['User Interface', 'Terminal (CLI) + Streamlit (Web)'],
        ['Cost', '$0 — all APIs have free tiers'],
    ]
)

doc.add_page_break()

# ============================================
# PAGE 2: FIRST APPROACH — 6 AGENTS
# ============================================
add_heading('First Approach — 6-Agent Architecture (Original Proposal)', level=0)
add_para('')

add_para(
    'This was the original architecture proposed in the dissertation. It uses '
    '6 AI agents, where each agent is powered by an LLM (Gemini). Every agent '
    'thinks, reasons, and uses tools to fetch data.'
)

add_bold_para('The 6 Agents and Their Jobs')

add_table(
    ['Agent Name', 'What It Does'],
    [
        ['1. Conversational Agent', 'Asks the user 8 questions one-by-one (destination, dates, budget, interests, etc.) until all information is collected'],
        ['2. Preferences Extractor', 'Reads the conversation and extracts structured data like origin, destination, dates, budget as JSON'],
        ['3. Flight Search Agent', 'Has a tool to search flights. It THINKS about what tool to use, CALLS the API, READS the result, and gives flight recommendations'],
        ['4. Hotel Search Agent', 'Same pattern — thinks, picks a hotel search tool, calls Booking.com API, reads results, recommends hotels'],
        ['5. Attractions Agent', 'Searches for attractions and restaurants using Serper web search, same think-call-read pattern'],
        ['6. Itinerary Coordinator', 'Reads all the data from above 5 agents and assembles a complete day-by-day itinerary'],
    ]
)

add_para('')
add_bold_para('How It Works (Step by Step)')
add_para('1. User types a request like "Plan a trip to Istanbul"')
add_para('2. Conversational Agent asks 8 questions (where from, dates, budget, etc.) — each question is 1 LLM call')
add_para('3. Preferences Extractor reads the conversation and converts it to JSON — 1 LLM call')
add_para('4. Flight Search Agent uses a CrewAI tool to call the fly-scraper API. The LLM must decide which tool to use, call it, then read the result — this takes multiple LLM calls due to the ReAct (Reasoning + Acting) loop')
add_para('5. Hotel Search Agent does the same for hotels — another ReAct loop with multiple LLM calls')
add_para('6. Attractions Agent does the same for attractions and restaurants — another ReAct loop')
add_para('7. Itinerary Coordinator reads all outputs and writes the final itinerary — 1 LLM call')

add_para('')
add_bold_para('The Problem with This Approach')
add_bullet('Each search agent wastes LLM calls just deciding which tool to use and reading the output — 5 LLM calls for extraction + 3 search agents + coordinator')
add_bullet('The ReAct loops cause errors — LLMs often mis-parse tool outputs and retry')
add_bullet('A simple task like "search for hotels" takes 30-60 seconds because the LLM is thinking instead of just calling the API')
add_bullet('Total: 5 LLM calls, 170-336 seconds per trip. With conversation included (8 more calls): 13 LLM calls')
add_bullet('The ablation study skips the conversational agent for fair comparison — both architectures benchmark against the same extraction + data + coordinator flow')

doc.add_page_break()

# ============================================
# PAGE 3: A2A PROTOCOL
# ============================================
add_heading('A2A Protocol — Agent-to-Agent Communication', level=0)
add_para('')
add_para(
    'A2A (Agent-to-Agent) is a communication protocol that lets agents send '
    'structured messages to each other. Think of it like email between agents — '
    'each message has a sender, receiver, type, and content.'
)

add_bold_para('How A2A Works')
add_bullet('8 Agent Cards are registered: conversational_agent, preferences_extractor, itinerary_coordinator, flight_data_provider, hotel_data_provider, attraction_data_provider, restaurant_data_provider, and user')
add_bullet('Each card defines: who it can send to, who it can receive from, what data it expects, and what data it returns')
add_bullet('When Agent A wants to send data to Agent B, the protocol checks if that communication is allowed by the registry')
add_bullet('Every trip generates 6 A2A messages — this message flow is displayed at the end for the dissertation')

add_para('')
add_bold_para('Example A2A Message Flow for One Trip:')
add_table(
    ['#', 'From', 'To', 'Type', 'Content'],
    [
        ['1', 'preferences_extractor', 'itinerary_coordinator', 'REQUEST', 'Extracted JSON preferences'],
        ['2', 'flight_data_provider', 'itinerary_coordinator', 'RESPONSE', 'Flight search results'],
        ['3', 'hotel_data_provider', 'itinerary_coordinator', 'RESPONSE', 'Hotel search results'],
        ['4', 'attraction_data_provider', 'itinerary_coordinator', 'RESPONSE', 'Attraction search results'],
        ['5', 'restaurant_data_provider', 'itinerary_coordinator', 'RESPONSE', 'Restaurant search results'],
        ['6', 'itinerary_coordinator', 'user', 'RESPONSE', 'Final itinerary'],
    ]
)

add_para('')
add_para(
    'Important: The A2A protocol is the SAME in both approaches (6-agent and 3-agent). '
    'The message flow, registry, and validation remain identical. Only the data-fetching '
    'layer changes — whether an LLM agent or a direct Python function sends the message.'
)

doc.add_page_break()

# ============================================
# PAGE 4: MCP
# ============================================
add_heading('MCP — Model Context Protocol', level=0)
add_para('')
add_para(
    'MCP (Model Context Protocol) is a standard way to give AI agents access to '
    'external tools like APIs. Think of it as "USB-C for AI" — a universal connector '
    'that lets any AI model use any tool.'
)

add_bold_para('How MCP is Set Up in This Project')
add_bullet('An MCP Server runs at src/server/mcp_server.py — it listens for tool requests over stdio (standard input/output)')
add_bullet('Tools are registered: search flights, search hotels, search attractions, search restaurants, calculator, web search')
add_bullet('The MCP server communicates using JSON-RPC format: a request comes in, the server does the work, and sends back the result')
add_bullet('MCP Tool wrappers are at src/tools/mcp_tools.py — these wrap the MCP server calls into LangChain StructuredTool format that CrewAI agents can use')

add_para('')
add_bold_para('MCP in the 6-Agent Approach')
add_bullet('Search agents (Flight, Hotel, Attraction) use MCP tools via CrewAI')
add_bullet('The agent says "I need to use the search_flights tool", CrewAI calls the MCP tool wrapper, which sends a JSON-RPC request to the MCP server, which calls the actual API, and returns the result')
add_bullet('This is slow because the LLM has to think about which tool to use at each step')

add_para('')
add_bold_para('MCP in the 3-Agent Approach')
add_bullet('The MCP server still exists and works — but the production flow BYPASSES it')
add_bullet('Instead of going through the MCP protocol layer, the code calls the Python functions directly (e.g., _call_fly_scraper_api())')
add_bullet('The MCP server is kept for the baseline comparison and for the coordinator agents tool access')

doc.add_page_break()

# ============================================
# PAGE 5: SECOND APPROACH — 3 AGENTS
# ============================================
add_heading('Second Approach — 3-Agent + Direct API (Optimized)', level=0)
add_para('')
add_para(
    'This is the improved architecture. The key insight: data fetching (calling APIs) '
    'does not need an LLM. A 3-line Python function can call an API faster and more '
    'reliably than an AI agent thinking about it. So we removed the 3 search agents '
    'and replaced them with direct Python function calls.'
)

add_bold_para('The 3 Agents in the Optimized Demo')
add_table(
    ['Agent Name', 'What It Does', 'LLM Calls'],
    [
        ['1. Preferences Extractor', 'Reads the user request and extracts structured JSON (origin, destination, dates, budget, interests)', '1 call'],
        ['2. Itinerary Coordinator', 'Takes all the data (flights, hotels, attractions, restaurants) and assembles a complete day-by-day itinerary', '1 call'],
    ]
)

add_para('')
add_bold_para('Note: The production CLI (run_cli.py) also has a Conversational Agent that asks 8 questions before extraction — making it 3 agents in the full pipeline. The optimized demo skips conversation for direct extraction.')

add_para('')
add_bold_para('What Changed?')

add_table(
    ['Before (6-Agent)', 'After (3-Agent)'],
    [
        ['Flight Search Agent (LLM decides tool, calls API, parses result)', '_call_fly_scraper_api() — direct Python function (no LLM)'],
        ['Hotel Search Agent (LLM decides tool, calls API, parses result)', 'search_hotels_comprehensive() — direct Python function (no LLM)'],
        ['Attractions Agent (LLM decides tool, calls API, parses result)', 'search_attractions() — direct Python function (no LLM)'],
        ['Conversational Agent (8 questions, 8 LLM calls)', 'Removed in optimized demo — goes straight to extraction'],
        ['5+ LLM calls per trip', '2 LLM calls per trip'],
        ['170-336 seconds per trip', '~33 seconds per trip'],
        ['Frequent parse errors in ReAct loops', 'No parse errors — Python handles data directly'],
    ]
)

add_para('')
add_bold_para('How It Works (Step by Step)')
add_para('1. User request goes directly to Preferences Extractor — 1 LLM call extracts JSON')
add_para('2. Python code calls four APIs directly in sequence (no LLM involved):')
add_bullet('_call_fly_scraper_api() — gets flight data')
add_bullet('search_hotels_comprehensive() — gets hotel data')
add_bullet('search_attractions() — gets attraction data')
add_bullet('search_restaurants() — gets restaurant data')
add_para('3. All data is combined and sent to Itinerary Coordinator — 1 LLM call to assemble the final itinerary')
add_para('4. A2A messages are still sent between all providers and coordinator (same protocol, same registry)')

doc.add_page_break()

# ============================================
# PAGE 6: COMPARISON & BENEFITS
# ============================================
add_heading('Comparison — Why the 3-Agent Approach is Better', level=0)
add_para('')

add_table(
    ['Metric', '6-Agent (Proposal)', '3-Agent (Optimized)', 'Improvement'],
    [
        ['LLM calls per trip', '5', '2', '60% fewer'],
        ['Average time per trip', '~230 seconds', '~33 seconds', '85% faster'],
        ['Parse errors', 'Common (ReAct loops)', 'None (direct calls)', 'More reliable'],
        ['A2A message flow', '6 messages', '6 messages', 'Same — protocol unchanged'],
        ['MCP Server', 'Used by all agents', 'Still exists, bypassed for speed', 'Available if needed'],
        ['Code complexity', 'Complex (agent reasoning)', 'Simple (function calls)', 'Easier to maintain'],
        ['API cost per trip', '~$0.011', '~$0.006', '~45% cheaper'],
    ]
)

add_para('')
add_bold_para('The Research Insight')
add_para(
    'The A2A multi-agent protocol is valuable — it provides clear message tracing, '
    'permission validation, and structured communication between components. But '
    'putting LLMs in the loop for deterministic API calls is wasteful. '
    'The optimized approach keeps the A2A protocol visible for the dissertation '
    'while removing LLM overhead from data fetching.'
)

add_para('')
add_bold_para('What This Proves')
add_bullet('Multi-agent architectures can be decoupled from the execution layer')
add_bullet('LLMs should only be used where reasoning adds value (understanding requests, assembling itineraries)')
add_bullet('Deterministic tasks (API calls) should use deterministic code (Python functions)')
add_bullet('The same A2A protocol works with both architectures — proving the protocol is independent of how data is fetched')

doc.add_page_break()

# ============================================
# PAGE 7: FILE STRUCTURE
# ============================================
add_heading('Project File Structure — What Each File Does', level=0)
add_para('')

add_table(
    ['File', 'Purpose'],
    [
        ['run_cli.py', 'PRODUCTION CLI — Interactive Q&A. Prompts user for trip request, calls orchestrator, prints itinerary.'],
        ['run_web.py', 'PRODUCTION WEB — Streamlit web interface for trip planning.'],
        ['demo_6agent_explained.py', 'EDUCATIONAL DEMO — Explains 6-agent architecture with MCP & A2A protocol details at each step. Auto-run.'],
        ['demo_3agent_explained.py', 'EDUCATIONAL DEMO — Explains 3-agent + Direct API architecture with MCP & A2A details. Auto-run.'],
        ['demo_comparison.py', 'COMPARISON — Runs both architectures back-to-back on same input. Shows side-by-side metrics table.'],
        ['run_6agent.py', 'EXECUTION — Runs 5 agents one-by-one showing each agent call and raw output live.'],
        ['run_3agent.py', 'EXECUTION — Runs extraction, 4 direct API calls (showing each response), then coordinator.'],
        ['generate_docx.py', 'Generates this DOCX document. Run: python generate_docx.py'],
        ['src/agents.py', 'Defines the 3 AI agents: conversational_agent (asks questions), preferences_extractor (parses JSON), itinerary_coordinator (assembles itinerary).'],
        ['src/tasks.py', 'Defines the 3 tasks that the agents perform. Each task has a description and expected output format.'],
        ['src/orchestrator.py', 'The main engine. Contains plan_trip() and plan_trip_from_transcript() which control the entire flow: conversation → extraction → API calls → coordination.'],
        ['src/comms/registry.py', 'A2A Agent Registry — defines 8 agent cards with their communication permissions (who can send to whom).'],
        ['src/comms/protocol.py', 'A2A Protocol implementation — message creation, validation, sending, history tracking.'],
        ['src/tools/mcp_tools.py', 'API wrapper functions that call external APIs (fly-scraper for flights, Booking.com for hotels, Serper for web search). Also wraps them as LangChain tools.'],
        ['src/tools/__init__.py', 'Exports all tools so agents can import them easily.'],
        ['src/server/mcp_server.py', 'MCP Server implementation. Runs over stdio, listens for JSON-RPC requests, calls APIs, returns results. 1155 lines.'],
        ['src/core/validators.py', 'Validates that the itinerary has the correct number of days and adds notices if days are missing.'],
        ['src/core/cache.py', 'Caches API responses to avoid duplicate calls.'],
        ['src/core/resilience.py', 'Retry logic for failed API calls.'],
        ['src/ui/app.py', 'Streamlit web app pages — chat input, itinerary display, comparison view.'],
        ['comparison/architecture_6agent.py', '6-agent baseline architecture for ablation study. Has 6 agents with MCP tools.'],
        ['comparison/architecture_3agent.py', '3-agent optimized architecture wrapper for comparison. Uses direct API calls.'],
        ['comparison/scenarios.py', '20 test scenarios with different destinations, budgets, trip lengths.'],
        ['comparison/run_comparison.py', 'Runs all 20 scenarios on both architectures, aggregates metrics, saves to JSON.'],
        ['.env', 'API keys file (GOOGLE_API_KEY, RAPIDAPI_KEY, SERPER_API_KEY, GEMINI_API_KEY). Not committed to git.'],
        ['AGENTS.md', 'Session tracking document — what was done, current state, next steps.'],
        ['pyproject.toml', 'Project dependencies — CrewAI 0.86.0, LiteLLM, Streamlit, etc.'],
    ]
)

# Save
output_path = os.path.join(os.path.dirname(__file__), 'Dissertation_Project_Explanation.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
