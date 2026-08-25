"""
Document plumbing: styles, headings, tables, citations and word counting.

Also holds val(), which every number in the dissertation passes through. It
records where each figure was printed, which is what the perturbation proof
checks.
"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from report.build.references import REFERENCES

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
FIGURES = os.path.join(ROOT, "report", "figures")

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x52, 0x51, 0x4E)

BODY_PT = 11
LINE_SPACING = 1.5

# The brief: "This assessment will be marked anonymously and should show your
# student number only." The name is deliberately absent from every generated
# artefact.
STUDENT_NUMBER = "25182589"
MODULE_CODE = "CMP7200"
MODULE_TITLE = "Individual Master's Project"
AWARD = "MSc Artificial Intelligence"
UNIVERSITY = "Birmingham City University"
FACULTY = "Faculty of Computing, Engineering and the Built Environment"
TITLE = ("Fewer Agents, Measured: Protocol-Mediated Multi-Agent Travel Planning "
         "and the Cost of Delegating Retrieval to a Language Model")
SUBMISSION_DATE = "7 September 2026"

# Words banned from the prose because they are filler that reads as generated
# text. Enforced by build_report.py, not by good intentions.
BANNED_WORDS = [
    "delve", "leverage", "leveraging", "robust", "seamless", "seamlessly",
    "cutting-edge", "moreover", "furthermore", "additionally",
]


class BuildError(RuntimeError):
    """A defect that must stop the build rather than reach the submission."""


# --------------------------------------------------------------------- citations
#   [@key]          -> (Author, Year)
#   [@key|n]        -> Author (Year)          narrative form
#   [@key1; @key2]  -> (Author1, Year1; Author2, Year2)
_CITE_RE = re.compile(r"\[@([a-z0-9]+)(\|n)?((?:\s*;\s*@[a-z0-9]+)*)\]")


@dataclass
class CitationLog:
    used: Dict[str, int] = field(default_factory=dict)

    def note(self, key: str) -> None:
        self.used[key] = self.used.get(key, 0) + 1

    def keys(self) -> List[str]:
        return sorted(self.used)

    def verify(self) -> None:
        unknown = [k for k in self.used if k not in REFERENCES]
        if unknown:
            raise BuildError(f"cited but not in references.py: {unknown}")
        uncited = [k for k in REFERENCES if k not in self.used]
        if uncited:
            raise BuildError(
                f"defined in references.py but never cited: {uncited}. "
                f"An uncited reference reads as padding — cite it or remove it.")
        missing_locator = [k for k in self.used if not REFERENCES[k].get("locator")]
        if missing_locator:
            raise BuildError(f"references without a DOI/arXiv id/URL: {missing_locator}")


CITATIONS = CitationLog()


def expand_citations(text: str) -> str:
    """Turn `[@key]` and `[@key|n]` into Harvard in-text citations."""
    def replace(match: re.Match) -> str:
        first, narrative, rest = match.group(1), match.group(2), match.group(3) or ""
        keys = [first] + re.findall(r"@([a-z0-9]+)", rest)
        for key in keys:
            if key not in REFERENCES:
                raise BuildError(f"unknown citation key {key!r}")
            CITATIONS.note(key)
        if narrative:
            if len(keys) > 1:
                raise BuildError(
                    f"narrative form takes one key, got {keys}; a narrative "
                    f"citation names its author in the sentence")
            ref = REFERENCES[keys[0]]
            return f"{ref['short']} ({ref['year']})"
        inner = "; ".join(f"{REFERENCES[k]['short']}, {REFERENCES[k]['year']}"
                          for k in keys)
        return f"({inner})"
    return _CITE_RE.sub(replace, text)


# ------------------------------------------------------------- measured values
@dataclass
class ValueLog:
    """
    Every measured value written into the document.

    Recorded so the perturbation check can assert that changing the source data
    changes the document. Without this the "no hardcoded numbers" claim is
    unfalsifiable.
    """
    entries: List[Tuple[str, str]] = field(default_factory=list)

    def record(self, source: str, rendered: str) -> None:
        self.entries.append((source, rendered))

    def rendered(self) -> List[str]:
        return [r for _, r in self.entries]

    def by_location(self) -> Dict[str, List[str]]:
        """Every rendered value grouped by the source line that produced it."""
        out: Dict[str, List[str]] = {}
        for source, value in self.entries:
            out.setdefault(source, []).append(value)
        return out


VALUES = ValueLog()


def val(value: Any, fmt: str = "{:,.0f}", *, source: str = "") -> str:
    """
    Format a measured value and record where in the source it came from.

    The call site is captured automatically. That is what lets the perturbation
    check compare the two builds POSITION BY POSITION rather than as two sets of
    strings: the same file and line means the same source expression, so a value
    that is identical across a corrupted and an uncorrupted build at the same
    location is genuinely suspicious. Comparing sets instead produced three false
    accusations of hardcoding in a row, each one a different number coinciding
    with an unrelated value elsewhere in the document.
    """
    try:
        rendered = fmt.format(value)
    except (ValueError, TypeError) as exc:
        raise BuildError(f"cannot format {value!r} with {fmt!r}: {exc}") from exc
    if not source:
        import sys as _sys
        frame = _sys._getframe(1)
        source = f"{os.path.basename(frame.f_code.co_filename)}:{frame.f_lineno}"
    VALUES.record(source, rendered)
    return rendered


# ------------------------------------------------------------------- the report
@dataclass
class Report:
    """A thin wrapper over python-docx that knows about this document's rules."""

    doc: Document = field(default_factory=Document)
    chapter: int = 0
    _figure_n: int = 0
    _table_n: int = 0
    _in_body: bool = False
    body_words: int = 0
    excluded_words: int = 0
    words_by_chapter: Dict[str, int] = field(default_factory=dict)
    _current: str = "front matter"
    figure_index: List[Tuple[str, str]] = field(default_factory=list)
    table_index: List[Tuple[str, str]] = field(default_factory=list)
    prose_blocks: List[Tuple[str, str]] = field(default_factory=list)
    all_text: List[str] = field(default_factory=list)

    # ------------------------------------------------------------- setup
    def __post_init__(self) -> None:
        style = self.doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(BODY_PT)
        fmt = style.paragraph_format
        fmt.line_spacing = LINE_SPACING
        fmt.space_after = Pt(8)
        # Match the East Asian font too, or Word silently substitutes for some
        # glyphs and the body text stops being uniformly 11 pt.
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), "Calibri")

        for section in self.doc.sections:
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)

        self._add_page_numbers()

    def _add_page_numbers(self) -> None:
        """
        Centred page numbers via a PAGE field.

        python-docx has no page-number API, so the field is inserted as raw
        WordprocessingML. The brief requires page numbers; a footer that says
        "Page" with no field would satisfy nothing.
        """
        footer = self.doc.sections[0].footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(end)
        run.font.size = Pt(9)
        run.font.color.rgb = MUTED

    # -------------------------------------------------------- word accounting
    def start_body(self, name: str) -> None:
        """Begin a section whose words count against the 12,000 limit."""
        self._in_body = True
        self._current = name
        self.words_by_chapter.setdefault(name, 0)

    def start_excluded(self, name: str) -> None:
        """Begin a section the brief excludes from the word count."""
        self._in_body = False
        self._current = name

    def _count(self, text: str) -> None:
        n = len(text.split())
        # Kept whole so cross-references can be resolved against the finished
        # document rather than against one chapter at a time.
        self.all_text.append(text)
        if self._in_body:
            self.body_words += n
            self.words_by_chapter[self._current] = \
                self.words_by_chapter.get(self._current, 0) + n
        else:
            self.excluded_words += n

    # ---------------------------------------------------------- text elements
    def h1(self, text: str, *, numbered: bool = True) -> None:
        if numbered:
            self.chapter += 1
            self._figure_n = 0
            self._table_n = 0
            text = f"{self.chapter}. {text}"
        self.doc.add_page_break()
        heading = self.doc.add_heading(text, level=1)
        heading.paragraph_format.space_after = Pt(10)
        self._count(text)

    def unnumbered_h1(self, text: str, *, page_break: bool = True) -> None:
        """
        A top-level heading outside the numbered chapters — Abstract, References,
        an appendix. Routed through here rather than straight to python-docx so
        the text still reaches the accumulator that cross-reference checking
        reads; appendix headings added directly were invisible to it, and every
        "Appendix X" reference was reported as dangling.
        """
        if page_break:
            self.doc.add_page_break()
        self.doc.add_heading(text, level=1)
        self._count(text)

    def h2(self, text: str) -> None:
        self.doc.add_heading(text, level=2)
        self._count(text)

    def h3(self, text: str) -> None:
        self.doc.add_heading(text, level=3)
        self._count(text)

    def p(self, text: str) -> None:
        """A body paragraph. Citations are expanded; words are counted."""
        expanded = expand_citations(" ".join(text.split()))
        paragraph = self.doc.add_paragraph(expanded)
        paragraph.paragraph_format.line_spacing = LINE_SPACING
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self._count(expanded)
        if self._in_body:
            self.prose_blocks.append((self._current, expanded))

    def bullets(self, items: Iterable[str]) -> None:
        for item in items:
            expanded = expand_citations(" ".join(item.split()))
            paragraph = self.doc.add_paragraph(expanded, style="List Bullet")
            paragraph.paragraph_format.line_spacing = LINE_SPACING
            self._count(expanded)

    def numbered_list(self, items: Iterable[str]) -> None:
        for item in items:
            expanded = expand_citations(" ".join(item.split()))
            paragraph = self.doc.add_paragraph(expanded, style="List Number")
            paragraph.paragraph_format.line_spacing = LINE_SPACING
            self._count(expanded)

    def quote(self, text: str) -> None:
        expanded = expand_citations(" ".join(text.split()))
        paragraph = self.doc.add_paragraph(expanded)
        paragraph.paragraph_format.left_indent = Inches(0.4)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.italic = True
            run.font.size = Pt(10)
        self._count(expanded)

    def code(self, text: str) -> None:
        """A verbatim block: log output, a schema, a traceback."""
        paragraph = self.doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(10)
        run = paragraph.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        rpr = run._element.get_or_add_rPr()
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), "Consolas")
        rfonts.set(qn("w:hAnsi"), "Consolas")
        rpr.append(rfonts)
        self._count(text)

    # --------------------------------------------------------------- captions
    def caption(self, text: str) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(expand_citations(" ".join(text.split())))
        run.italic = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = MUTED
        self._count(text)

    def figure(self, relative_path: str, caption: str, *, width: float = 6.3) -> str:
        """
        Insert a generated figure with an auto-numbered caption.

        A missing figure raises. Silently printing "[figure missing]" into a
        submitted document is worse than failing the build.
        """
        path = os.path.join(FIGURES, relative_path)
        if not os.path.exists(path):
            raise BuildError(
                f"figure {relative_path} does not exist. Run "
                f"report/build/make_diagrams.py and report/build/make_charts.py first.")
        self._figure_n += 1
        tag = f"Figure {self.chapter}.{self._figure_n}"
        self.doc.add_picture(path, width=Inches(width))
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.caption(f"{tag}: {caption}")
        self.figure_index.append((tag, caption))
        return tag

    def table(self, headers: Sequence[str], rows: Sequence[Sequence[Any]],
              caption: str, *, widths: Optional[Sequence[float]] = None,
              font_pt: float = 9.5) -> str:
        self._table_n += 1
        tag = f"Table {self.chapter}.{self._table_n}"
        self.caption(f"{tag}: {caption}")
        table = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, head in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(head))
            run.bold = True
            run.font.size = Pt(font_pt)
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0
            self._count(str(head))

        for r, row in enumerate(rows, start=1):
            for c, value in enumerate(row):
                cell = table.rows[r].cells[c]
                cell.text = ""
                run = cell.paragraphs[0].add_run(expand_citations(str(value)))
                run.font.size = Pt(font_pt)
                cell.paragraphs[0].paragraph_format.line_spacing = 1.0
                self._count(str(value))

        if widths:
            for row in table.rows:
                for cell, w in zip(row.cells, widths):
                    cell.width = Inches(w)
        self.doc.add_paragraph()
        self.table_index.append((tag, caption))
        return tag

    # ----------------------------------------------------------------- output
    def save(self, path: str) -> str:
        os.makedirs(os.path.dirname(path), exist_ok=True)
        self.doc.save(path)
        return path


# --------------------------------------------------------------- text checks
def find_banned_words(blocks: Sequence[Tuple[str, str]]) -> List[Tuple[str, str]]:
    """Banned filler words, with the chapter they appear in."""
    hits: List[Tuple[str, str]] = []
    for chapter, text in blocks:
        lowered = text.lower()
        for word in BANNED_WORDS:
            if re.search(rf"\b{re.escape(word)}\b", lowered):
                hits.append((chapter, word))
    return hits


# Proper nouns whose official spelling is American. These are names, not prose,
# so changing them would misquote the source.
_SPELLING_EXCEPTIONS = ("World Tourism Organization",)

# US spellings that must not appear, with the British form used instead. The
# document commits to British spelling; mixing the two reads as carelessness and
# is the kind of thing a marker notices on the first page.
_US_SPELLINGS = [
    ("optimize", "optimise"), ("optimizing", "optimising"),
    ("optimized", "optimised"), ("optimization", "optimisation"),
    ("analyze", "analyse"), ("analyzed", "analysed"),
    ("behavior", "behaviour"), ("color", "colour"),
    ("organization", "organisation"), ("recognize", "recognise"),
    ("modeled", "modelled"), ("labeled", "labelled"),
    ("traveler", "traveller"), ("travelers", "travellers"),
    ("center", "centre"), ("defense", "defence"),
    ("summarize", "summarise"), ("normalize", "normalise"),
    ("fulfill", "fulfil"), ("practise", "practice"),
]


def find_spelling_inconsistencies(text: str) -> List[Tuple[str, str]]:
    """US spellings in a document that commits to British English."""
    scrubbed = text
    for phrase in _SPELLING_EXCEPTIONS:
        scrubbed = scrubbed.replace(phrase, "")
    hits = []
    for us, uk in _US_SPELLINGS:
        if re.search(rf"\b{us}\b", scrubbed, re.IGNORECASE):
            hits.append((us, uk))
    return hits


def find_dangling_references(text: str) -> List[str]:
    """
    Cross-references pointing at something the document does not contain.

    A reference to a section that was renumbered, or an appendix that was
    reordered, sends a marker to the wrong page. Both happened during drafting —
    appendices I and J were once emitted out of order — so both are checked here
    rather than trusted.
    """
    problems = []

    referenced = set(re.findall(r"Appendix ([A-Z])\b", text))
    present = set(re.findall(r"Appendix ([A-Z]) —", text))
    for letter in sorted(referenced - present):
        problems.append(f"Appendix {letter} is referenced but does not exist")

    sections = set(re.findall(r"Section (\d+\.\d+)", text))
    headings = set(re.findall(r"(?:^|\n)(\d+\.\d+)\s", text))
    for ref in sorted(sections - headings):
        problems.append(f"Section {ref} is referenced but has no heading")

    return problems


def _shingles(text: str, n: int = 8) -> set:
    words = re.findall(r"[a-z0-9']+", text.lower())
    return {" ".join(words[i:i + n]) for i in range(max(0, len(words) - n + 1))}


def find_duplicate_prose(blocks: Sequence[Tuple[str, str]],
                         threshold: float = 0.35) -> List[Tuple[str, str, float]]:
    """
    Paragraph pairs sharing more than `threshold` of their 8-word sequences.

    Repeating an argument across two chapters wastes words against a hard limit
    and reads as padding. An 8-word shingle is long enough that shared technical
    vocabulary does not trigger it, and short enough to catch a reused sentence.
    """
    prepared = [(chapter, text, _shingles(text)) for chapter, text in blocks
                if len(text.split()) >= 25]
    flagged: List[Tuple[str, str, float]] = []
    for i, (ch_a, text_a, sh_a) in enumerate(prepared):
        for ch_b, text_b, sh_b in prepared[i + 1:]:
            if not sh_a or not sh_b:
                continue
            overlap = len(sh_a & sh_b) / min(len(sh_a), len(sh_b))
            if overlap > threshold:
                flagged.append((f"{ch_a}: {text_a[:70]}...",
                                f"{ch_b}: {text_b[:70]}...", round(overlap, 3)))
    return flagged
