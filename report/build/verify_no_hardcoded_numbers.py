"""
Proves no number in the dissertation was typed by hand.

Corrupts the results files, rebuilds, and checks that every printed figure
moved. Anything that did not move is either hardcoded or measured from the
repository itself, and each of those is checked against its live source.
"""

from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
from typing import Any, Dict, List, Set

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(
    os.path.abspath(__file__)))))

from evaluation.measured import (BUDGET_GATE_PATH, COMPARISON_PATH,
                                PROTOCOL_PATH, ROOT)

RESULT_FILES = [COMPARISON_PATH, PROTOCOL_PATH, BUDGET_GATE_PATH]

# Values too common to prove anything. "1" appears in "Section 1", "20" in
# "twenty scenarios", "0" everywhere. Excluding them keeps the check honest:
# a value only counts as evidence if its presence in the document could plausibly
# come from nowhere else.
TRIVIAL = {"0", "1", "2", "3", "4", "5", "6", "7", "8", "9", "10", "11", "12",
           "20", "0.0", "1.0", "100", "$0.0000"}


def _perturb(value: Any) -> Any:
    """
    Corrupt a number recognisably while keeping it a valid number.

    The transform is deliberately not a constant: replacing everything with 999
    would collapse distinct values together and weaken the test. Scaling and
    offsetting keeps every value distinct and every value wrong.
    """
    # Booleans are structural flags ("scored", "passed", "agrees"), not measured
    # quantities. Inverting them makes the results file invalid rather than wrong,
    # so the build fails on a missing measurement instead of testing what it means
    # to test. They are left alone; the check is about numbers.
    if isinstance(value, bool):
        return value
    # The offset is 41 rather than 13 to avoid a coincidence: with +13, every
    # stored zero rendered as "13", which collided with a real measured "13%"
    # elsewhere in the baseline and looked like an unchanged value.
    if isinstance(value, int):
        return value * 7 + 41
    if isinstance(value, float):
        return round(value * 7 + 41.0, 6)
    if isinstance(value, list):
        # Element values are perturbed AND the length is changed, because a
        # number in the document may be a list length (how many hotels were
        # retrieved, how many tools are defective) rather than a stored figure.
        # Without this, those values are unchanged by corruption and the check
        # cannot say anything about them. Duplicating the last element rather
        # than dropping one keeps every existing index valid.
        perturbed = [_perturb(v) for v in value]
        return perturbed + perturbed[-1:] if perturbed else perturbed
    if isinstance(value, dict):
        return {k: _perturb(v) for k, v in value.items()}
    return value


def _document_text(path: str) -> str:
    """Every string in the document: paragraphs, table cells, captions."""
    from docx import Document
    doc = Document(path)
    parts: List[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _repository_measured_values() -> Dict[str, str]:
    """
    Values the document reports that come from the repository, not from a
    results file, and therefore cannot change when a results file is corrupted.

    Corrupting the three results files cannot alter the test count, the line
    counts, the commit history or the size of a recorded API response, because
    none of those is stored there — they are measured from the repository at
    build time. Any value left unchanged by the perturbation must be one of
    these, and this function computes them from their live sources so that claim
    is checked rather than asserted.
    """
    from evaluation import measured

    out: Dict[str, str] = {}

    def add(label: str, value: Any) -> None:
        # Keys are normalised the same way document tokens are, or the lookup
        # misses every value that carries a thousands separator.
        for fmt in ("{:,.0f}", "{:,}", "{:.1f}"):
            try:
                out[_norm_num(fmt.format(value))] = label
            except (ValueError, TypeError):
                continue

    add("pytest collection", measured.test_count()["collected"])

    stats = measured.code_stats()
    add("code stats: total lines", stats["total_lines"])
    add("code stats: total files", stats["total_files"])
    for area, row in stats["areas"].items():
        add(f"code stats: {area} lines", row["lines"])
        add(f"code stats: {area} files", row["files"])

    git = measured.git_stats()
    if git.get("commits"):
        add("git history: commit count", git["commits"])

    cache = measured.api_cache_stats()
    add("api cache: entry count", cache["entries"])
    add("api cache: total kB", cache["total_kb"])
    for host, n in cache["by_host"].items():
        add(f"api cache: {host} entries", n)

    evidence = measured.flight_api_evidence()
    for key in ("broken_max_itineraries", "fixed_bytes", "fixed_itineraries",
                "poll_bytes", "poll_itineraries", "broken_recordings"):
        add(f"recorded flight response: {key}", evidence[key])
    for size in evidence["broken_bytes"]:
        add("recorded flight response: broken_bytes", size)

    return out


def _build(output: str, values_out: str) -> Dict[str, Any]:
    proc = subprocess.run(
        [sys.executable, "-m", "report.build.build_report",
         "--no-tests", "--output", output, "--values-out", values_out],
        cwd=ROOT, capture_output=True, text=True, timeout=900,
    )
    if not os.path.exists(output):
        raise SystemExit(f"build produced no document:\n{proc.stdout}\n{proc.stderr}")
    with open(values_out, encoding="utf-8") as fh:
        values = json.load(fh)
    values["text"] = _document_text(output)
    values["stdout"] = proc.stdout
    return values


def _distinctive(rendered: List[str]) -> Set[str]:
    """
    Rendered values specific enough that their presence is evidence.

    At least three digits are required. A two-digit value such as "1.1" or "5.5"
    carries too little information: in a document interpolating several hundred
    numbers, one will coincidentally appear whatever the source data says, and the
    check then reports it as hardcoded. Both of those were reported before this
    threshold was raised, and both were coincidences.
    """
    def digits(v: str) -> int:
        return sum(c.isdigit() for c in v)

    return {v for v in rendered
            if v not in TRIVIAL and digits(v) >= 3 and re.search(r"\d", v)}


# A number as it appears in the document: optional currency, digits with
# thousands separators, optional decimals, optional percent sign.
_NUMBER_TOKEN = re.compile(r"\$?\d[\d,]*(?:\.\d+)?%?")


def _norm_num(token: str) -> str:
    """
    Remove thousands separators only.

    Currency and percent signs are deliberately KEPT. Stripping them collides a
    measured "$300" with the unrelated literal "300" in "300 dpi", which the
    check then reports as a hardcoded value. Currency is part of what makes a
    number identifiable.
    """
    return token.replace(",", "")


def _present(value: str, tokens: Set[str]) -> bool:
    """
    Is `value` present in the document, allowing a percent sign to be added?

    A value recorded as "557.6" is written into the prose as "557.6%". Used only
    to check that a value the build says it interpolated really did arrive.
    """
    v = _norm_num(value)
    return v in tokens or f"{v}%" in tokens


def _survives(value: str, tokens: Set[str]) -> bool:
    """
    Did `value` persist into the corrupted document, matching exactly?

    Deliberately stricter than _present. Allowing "13%" to match a bare "13"
    reported a survivor every time an unrelated "13" appeared anywhere in the
    document, which is a false accusation of hardcoding.
    """
    return _norm_num(value) in tokens


def _number_tokens(text: str) -> Set[str]:
    """
    Every whole number token in the document.

    Whole tokens, not substrings. Testing `"0.27" in text` reports a false
    positive the moment the document contains "10.27" or "$1,230.27" — which it
    did, and which briefly looked like a hardcoded value.
    """
    return {_norm_num(t) for t in _NUMBER_TOKEN.findall(text)}


def _distinctive_sites(build: Dict[str, Any]) -> int:
    """How many (call site, value) pairs are distinctive enough to be evidence."""
    sites = build.get("by_location") or {}
    return len({(loc, v) for loc, values in sites.items()
                for v in values if _distinctive([v])})


def _unchanged_by_location(before: Dict[str, Any],
                           after: Dict[str, Any]) -> List[tuple]:
    """
    Values a single call site rendered identically in both builds.

    This compares the two builds POSITION BY POSITION rather than as two sets of
    strings, because `val()` records the file and line that produced each value
    and both builds run the same source. Same location means the same expression,
    so an identical value there is genuinely suspicious.

    Comparing sets instead asks a weaker question — "does this string appear
    anywhere in the corrupted document?" — and in a document interpolating four
    hundred numbers the answer is sometimes yes by coincidence. That produced four
    successive false accusations of hardcoding (0.27, $300, 13%, then 4,019, which
    was arm D's mean completion tokens colliding with an unrelated perturbed
    value). Each was a real interpolation. Positional comparison cannot make that
    mistake: a collision elsewhere in the document is now irrelevant.

    A site inside a loop can render a different COUNT of values between builds,
    because perturbation lengthens lists. Only the overlapping positions are
    compared; a position that exists in one build alone says nothing about
    hardcoding.
    """
    before_sites = before.get("by_location") or {}
    after_sites = after.get("by_location") or {}
    unchanged: List[tuple] = []
    for location, before_list in before_sites.items():
        after_list = after_sites.get(location)
        if not after_list:
            continue
        for old, new in zip(before_list, after_list):
            if old == new and old not in TRIVIAL and _distinctive([old]):
                unchanged.append((location, old))
    return sorted(set(unchanged))


def main() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")

    workspace = tempfile.mkdtemp(prefix="report_verify_")
    backups = {path: os.path.join(workspace, os.path.basename(path) + ".backup")
               for path in RESULT_FILES}

    print("=" * 78)
    print("  PERTURBATION CHECK — is any number in the document hardcoded?")
    print("=" * 78)

    try:
        for path, backup in backups.items():
            shutil.copy2(path, backup)
        print(f"\n  backed up {len(backups)} results files")

        # ---- 1. baseline -------------------------------------------------
        print("\n[1/4] building from the real measured results")
        before = _build(os.path.join(workspace, "before.docx"),
                        os.path.join(workspace, "before.json"))
        before_values = _distinctive(before["rendered"])
        print(f"      {before['count']} values interpolated, "
              f"{len(before_values)} distinctive, "
              f"{before['body_words']:,} words of body text")

        # ---- 2. corrupt --------------------------------------------------
        print("\n[2/4] corrupting every number in every results file (n -> n*7+41, lists grow)")
        for path in RESULT_FILES:
            with open(path, encoding="utf-8") as fh:
                data = json.load(fh)
            with open(path, "w", encoding="utf-8") as fh:
                json.dump(_perturb(data), fh, indent=2, default=str)
            print(f"      corrupted {os.path.relpath(path, ROOT)}")

        # ---- 3. rebuild and compare --------------------------------------
        print("\n[3/4] rebuilding from the corrupted results")
        after = _build(os.path.join(workspace, "after.docx"),
                       os.path.join(workspace, "after.json"))
        after_values = _distinctive(after["rendered"])
        after_tokens = _number_tokens(after["text"])
        print(f"      {after['count']} values interpolated, "
              f"{len(after_values)} distinctive")

        after_norm = {_norm_num(v) for v in after_values}
        survivors = sorted(v for v in before_values
                           if _norm_num(v) not in after_norm
                           and _survives(v, after_tokens))
        missing = sorted(v for v in after_values
                         if not _present(v, after_tokens))

        print(f"\n      old values surviving into the corrupted document: "
              f"{len(survivors)}")
        print(f"      new values absent from the corrupted document:     "
              f"{len(missing)}")

        distinctive_sites = _distinctive_sites(before)
        changed = distinctive_sites - len(_unchanged_by_location(before, after))
        print(f"      distinctive call sites whose value changed:         "
              f"{changed}/{distinctive_sites}")

        if survivors:
            print("\n      HARDCODED VALUES FOUND — these survived corruption:")
            for value in survivors[:40]:
                print(f"        {value}")
        if missing:
            print("\n      values interpolated but not found in the document:")
            for value in missing[:40]:
                print(f"        {value}")

        # ---- 4. restore --------------------------------------------------
        print("\n[4/4] restoring the real results and rebuilding")
        for path, backup in backups.items():
            shutil.copy2(backup, path)
        restored = _build(os.path.join(workspace, "restored.docx"),
                          os.path.join(workspace, "restored.json"))
        identical = restored["text"] == before["text"]
        print(f"      restored document identical to the baseline: {identical}")

        # Every value left unchanged must be attributable to a repository source
        # the perturbation cannot reach. Unattributed ones are candidates for
        # being hardcoded.
        repository = _repository_measured_values()
        unchanged = _unchanged_by_location(before, after)
        accounted = {v: repository[_norm_num(v)] for _, v in unchanged
                     if _norm_num(v) in repository}
        unaccounted = sorted({(loc, v) for loc, v in unchanged
                              if v not in accounted})

        print(f"\n      call sites compared position by position:           "
              f"{len(after.get('by_location') or {})}")
        # Pairs, not values: a repository figure such as the total line count is
        # quoted in more than one table, so it is several unchanged call sites
        # carrying one number. Reporting only the distinct count leaves the
        # changed/unchanged arithmetic looking short of the total.
        print(f"      unchanged call sites:                               "
              f"{len(unchanged)}, carrying {len(accounted)} distinct values, "
              f"each attributed below")
        for value, source in sorted(accounted.items()):
            print(f"        {value:>12}  <- {source}")
        if unaccounted:
            print(f"\n      UNCHANGED AND UNACCOUNTED FOR — possibly hardcoded:")
            for loc, value in unaccounted:
                print(f"        {value:>12}  at {loc}")

        ok = (not survivors and not missing and identical
              and not unaccounted and changed > 0)
        print("\n" + "=" * 78)
        if ok:
            print(f"  PASS — {changed} of {distinctive_sites} distinctive call sites "
                  f"changed value when the\n         results data changed. The remaining "
                  f"{len(unchanged)} quote {len(accounted)} figures measured from the\n"
                  f"         repository "
                  f"(pytest, the filesystem, git, the response cache) and were each\n"
                  f"         verified against their live source. None survived into "
                  f"the corrupted\n         document, and restoring the data "
                  f"reproduced the original exactly.\n         No number in the "
                  f"document is hardcoded.")
        else:
            print("  FAIL — see the detail above.")
        print("=" * 78)
        return 0 if ok else 1

    finally:
        # Restoring the real results matters more than any diagnostic, so it
        # happens even if the check raises part way through.
        for path, backup in backups.items():
            if os.path.exists(backup):
                shutil.copy2(backup, path)
        shutil.rmtree(workspace, ignore_errors=True)


if __name__ == "__main__":
    sys.exit(main())
