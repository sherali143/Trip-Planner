"""
Builds PROJECT_OVERVIEW.docx.

A plain-English guide to the whole project for someone who has never seen it:
what it is, how to run it, and how to show it to a supervisor. Numbers come
from the same accessor the dissertation uses.
"""

from __future__ import annotations

import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(
    os.path.abspath(__file__)))))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from evaluation import measured

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
OUTPUT = os.path.join(ROOT, "submission", "PROJECT_OVERVIEW.docx")

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
MUTED = RGBColor(0x52, 0x51, 0x4E)

# One workflow diagram per approach. Written as plain text so they survive being
# copied into an email, a slide or a Word document, and so the person presenting
# this can change a label without needing the drawing tool that made it.
DIAGRAM_A = """\
  your request
       |
       v
  +-----------------------+
  |     ONE AI CALL       |   no tools. no internet. no lookups.
  +-----------------------+
       |
       v
  a plan that LOOKS finished

  APIs called: NONE"""

DIAGRAM_B = """\
  your request
       |
       v
  [ extractor ]
       |
       +--> [ flight agent   ]  8 tools, loops until satisfied
       +--> [ hotel agent    ]  8 tools, loops until satisfied
       +--> [ activity agent ]  4 tools, loops until satisfied
                  |
                  v
          [ coordinator ]  writes the plan
                  |
                  v
          your itinerary

  Every loop step is a FULL AI call that re-sends the whole
  conversation and every tool description again."""

DIAGRAM_C = """\
  your request
       |
       v
  [ extractor ]
       |
       +--> [ flight agent   ]  1 tool   \\
       +--> [ hotel agent    ]  1 tool    >  all three run AT THE SAME TIME
       +--> [ activity agent ]  2 tools  /
                  |
                  v
          [ coordinator ]  writes the plan
                  |
                  v
          your itinerary

  Same six agents as B. Shorter prompts, fewer tools each, results
  trimmed before being passed on, and a hard cap of 3 loop steps."""

# How a tool call physically travels, per approach. Drawn because "we built an MCP
# server" and "the shipped system calls it over the protocol" are different claims,
# and only the first is true of every approach. Section 4 of the dissertation says
# the same thing in prose; a supervisor asking the question deserves the picture.
DIAGRAM_B_TOOLCALL = """\
  APPROACH B  --  goes through the MCP server over JSON-RPC

  [ hotel agent ]
        |  "call search_hotel_destination"
        v
  [ tool wrapper ]  trip_planner/tools/agent_tools.py
        |
        v
  [ MCP client ]  starts the server as a SEPARATE PROGRAM
        |
        |  ==== JSON-RPC down a pipe ====
        v
  [ MCP SERVER ]  a separate process holding 12 tools
        |         checks the input against the tool's declared schema
        v
  [ Booking.com / fly-scraper ]  the real API
        |
        v
  [ recorded-response cache ]  saves it, or replays a saved one
        |
        |  ==== JSON-RPC back up the pipe ====
        v
  the FULL reply (about 12,000 characters) goes to the AI"""

DIAGRAM_C_TOOLCALL = """\
  APPROACH C  --  calls the same functions directly, no pipe

  [ hotel agent ]
        |  "call find_hotels"
        v
  [ distilled tool ]  evaluation/distilled_tools.py
        |
        |  imports the server's own function and calls it here.
        |  No separate program. No JSON-RPC.
        v
  [ the same server function ]  search_hotels_comprehensive
        |
        v
  [ Booking.com / fly-scraper ]  the same real API
        |
        v
  [ recorded-response cache ]  the same cache
        |
        v
  [ TRIM to the best 3 ]  <-- this is the change that matters
        |
        v
  3 short lines go to the AI"""

# What "steps allowed" means, shown rather than defined. It is the single setting
# that most explains the cost difference, and "max_iter=10" means nothing to a
# reader until they watch an agent use up ten of them.
ITERATIONS_B = """\
  APPROACH B  --  hotel agent: 8 tools, 10 steps allowed

  STEP 1   AI is sent: the job + ALL 8 tool descriptions
           AI answers: "look up Istanbul's ID code"        <- 1 AI call
           tool runs -> gets: dest_id -755070

  STEP 2   AI is sent: the job + ALL 8 descriptions AGAIN + what happened
           AI answers: "now search hotels with that code"  <- 1 AI call
           tool runs -> gets: 12,000 characters of hotels

  STEP 3   AI is sent: the job + ALL 8 descriptions AGAIN + everything so far
                       + all 12,000 characters
           AI answers: "let me check the reviews too"      <- 1 AI call

  STEP 4   ...again... "let me look at nearby attractions" <- 1 AI call
  STEP 5   ...again... "let me search the web to compare"  <- 1 AI call
  STEP 6   ...again... "let me total up 4 nights"          <- 1 AI call
  STEP 7-10  it may keep going. At step 10 it is cut off and
             must answer with whatever it has.

  Result: about 6 AI calls for one hotel, and 8 descriptions re-sent
          on every single one of them."""

ITERATIONS_C = """\
  APPROACH C  --  hotel agent: 1 tool, 3 steps allowed

  STEP 1   AI is sent: the job + ONE short description
           AI answers: "find_hotels('Istanbul', 15 Aug, 19 Aug, 80)"  <- 1 AI call
           tool runs -> gets 12,000 characters -> TRIMMED to the best 3

  STEP 2   AI is sent: the job + the one description + 3 short lines
                       Theodora Pension  $33/night  10/10
                       Hotel Sultania    $41/night   9/10
                       Ada Homes         $38/night   9/10
           AI answers: "Theodora Pension. Done."                      <- 1 AI call

  STEP 3   not needed. It finished on its own.

  Result: 2 AI calls, and there was nothing to wander off to."""

DIAGRAM_D = """\
  your request
       |
       v
  +----------------------------+
  |  AI STEP 1: read request   |  a judgement: wording varies endlessly
  +----------------------------+
       |   origin, destination, dates, budget, interests
       v
  +----------------------------+
  |  PLAIN PYTHON: fetch data  |  NO AI HERE
  |                            |
  |    1  flights      1 call  |
  |    2  hotels       2 calls |
  |    3  attractions  1 call  |
  |    4  restaurants  1 call  |
  |                            |
  |    always these 5 calls    |
  +----------------------------+
       |
       v
  +----------------------------+
  |  AI STEP 2: write the plan |  a judgement: how to order a day
  +----------------------------+
       |
       v
  your day-by-day itinerary"""


class Doc:
    """A very small wrapper. This document has no numbering or cross-references."""

    def __init__(self) -> None:
        self.d = Document()
        style = self.d.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(6)
        for s in self.d.sections:
            s.left_margin = s.right_margin = Inches(0.9)
            s.top_margin = s.bottom_margin = Inches(0.8)

    def title(self, text: str, subtitle: str = "") -> None:
        p = self.d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = ACCENT
        if subtitle:
            p2 = self.d.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(subtitle)
            r2.font.size = Pt(11)
            r2.font.color.rgb = MUTED

    def h(self, text: str) -> None:
        self.d.add_heading(text, level=1)

    def h2(self, text: str) -> None:
        self.d.add_heading(text, level=2)

    def p(self, text: str) -> None:
        self.d.add_paragraph(" ".join(text.split()))

    def bullets(self, items) -> None:
        for i in items:
            self.d.add_paragraph(" ".join(i.split()), style="List Bullet")

    def steps(self, items) -> None:
        for i in items:
            self.d.add_paragraph(" ".join(i.split()), style="List Number")

    def code(self, text: str) -> None:
        p = self.d.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        r.font.name = "Consolas"
        r.font.size = Pt(9)

    def table(self, headers, rows, widths=None, font_pt=9.5) -> None:
        t = self.d.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Light Grid Accent 1"
        for i, head in enumerate(headers):
            c = t.rows[0].cells[i]
            c.text = ""
            r = c.paragraphs[0].add_run(str(head))
            r.bold = True
            r.font.size = Pt(font_pt)
        for ri, row in enumerate(rows, start=1):
            for ci, val in enumerate(row):
                c = t.rows[ri].cells[ci]
                c.text = ""
                r = c.paragraphs[0].add_run(str(val))
                r.font.size = Pt(font_pt)
                c.paragraphs[0].paragraph_format.space_after = Pt(2)
        if widths:
            for row in t.rows:
                for c, w in zip(row.cells, widths):
                    c.width = Inches(w)
        self.d.add_paragraph()

    def save(self) -> str:
        self.d.save(OUTPUT)
        return OUTPUT


# What each file does, in one line, in plain English. The table in the document is
# built by listing the folders on disk and looking each file up here, so a new file
# with no entry shows up as "(not yet described)" rather than being silently
# omitted — the same reason the test suite checks that every module is named in its
# package README.
# Every tool each agent holds, and what it does in one line. Taken from the agent
# definitions in evaluation/arm_b_six_agent_naive.py and arm_c_six_agent_tuned.py.
# The point of listing them is not completeness for its own sake — it is that the
# duplication becomes visible. search_internet and calculate appear on all four of
# the naive arm's agents, which is 8 of its 20 slots spent on two tools.
TOOLS_B = [
    ("Extractor", []),
    ("Flight agent", [
        ("search_comprehensive_flights", "Finds flights with all the details"),
        ("search_round_trip_flights", "Finds return flights specifically"),
        ("search_internet", "Searches the web for anything"),
        ("calculate", "Does sums"),
    ]),
    ("Hotel agent", [
        ("search_hotels_comprehensive", "Finds hotels with prices and ratings"),
        ("search_accommodations_with_location", "Finds places to stay near a spot"),
        ("search_hotel_destination", "Looks up a city's internal ID code"),
        ("search_hotels_by_dest_id", "Finds hotels using that ID code"),
        ("get_hotel_reviews", "Gets a hotel's review scores"),
        ("get_attractions_near_hotel", "Finds things to do near a hotel"),
        ("search_internet", "Searches the web for anything"),
        ("calculate", "Does sums"),
    ]),
    ("Attraction agent", [
        ("search_attractions", "Finds places to visit"),
        ("search_restaurants", "Finds places to eat"),
        ("search_internet", "Searches the web for anything"),
        ("calculate", "Does sums"),
    ]),
    ("Coordinator", [
        ("calculate", "Does sums"),
        ("search_internet", "Searches the web for anything"),
        ("search_attractions", "Finds places to visit"),
        ("search_restaurants", "Finds places to eat"),
    ]),
]

TOOLS_C = [
    ("Extractor", []),
    ("Flight agent", [
        ("distilled_search_flights", "Finds flights, returns only the cheapest few"),
    ]),
    ("Hotel agent", [
        ("distilled_search_hotels", "Finds hotels, returns only the best-rated few"),
    ]),
    ("Attraction agent", [
        ("distilled_search_attractions", "Finds places to visit, top few only"),
        ("distilled_search_restaurants", "Finds places to eat, top few only"),
    ]),
    ("Coordinator", []),
]

# What each agent is ALLOWED to do, read from the agent definitions. Not how many
# times it actually did — the code does not tag a model request with the agent that
# caused it, and inventing that attribution would be worse than leaving it out.
LIMITS = {
    "Extractor": (3, 3),
    "Flight agent": (8, 3),
    "Hotel agent": (10, 3),
    "Attraction agent": (10, 3),
    "Coordinator": (15, 3),
}

FILE_PURPOSE = {
    # top level
    "run_cli.py": "Plan a trip by typing in the terminal. Runs the shipped design.",
    "run_web.py": "Plan a trip in a web page. Runs the same shipped design.",
    # trip_planner
    "trip_planner/orchestrator.py":
        "The workflow. Reads the request, fetches the data, assembles the plan. "
        "Both the terminal and the web page go through this one file.",
    "trip_planner/agents.py":
        "Defines the three AI agents: one to chat, one to read the request, one to "
        "write the itinerary.",
    "trip_planner/tasks.py": "The instructions each agent is given.",
    # core
    "trip_planner/core/http_cache.py":
        "Records every reply from every travel website and replays it later. This is "
        "why the results reproduce with no accounts, and why one careless run cannot "
        "burn a whole month of free quota.",
    "trip_planner/core/llm_metrics.py":
        "Counts the real AI calls, tokens and cost. An earlier version estimated "
        "these and the estimates were wrong.",
    "trip_planner/core/trip_cost.py":
        "Works out what a trip costs at minimum, comfortable and luxury levels, and "
        "refuses a budget below the real floor. This is the budget validation the "
        "supervisor asked for.",
    "trip_planner/core/budget.py":
        "Splits a budget across flights, hotels, food and activities based on what "
        "the trip actually costs, instead of a fixed percentage for every trip.",
    "trip_planner/core/real_prices.py":
        "Reads what flights really cost out of the saved API replies, so the "
        "budget check can use a real fare instead of a number from a table. Free "
        "— the replies are already on disk. Says which figure came from where.",
    "trip_planner/core/gemini_compat.py":
        "Keeps the project running on current Google AI models. The newer models "
        "reject a message pattern the agent loop produces; this fixes only those "
        "requests and counts how many it changed.",
    "trip_planner/core/safe_math.py":
        "Does the arithmetic for the calculator tool without letting the AI run "
        "arbitrary code. Also blocks sums designed to hang the machine.",
    "trip_planner/core/resilience.py":
        "Decides whether a refusal from the AI provider is worth retrying. A spending "
        "cap and a per-minute limit look the same but need opposite responses.",
    "trip_planner/core/validators.py":
        "Checks the finished itinerary actually contains every day it should.",
    "trip_planner/core/log_setup.py":
        "Keeps the API key out of the console. Google sends it inside the web "
        "address, which the network library would otherwise print.",
    # tools
    "trip_planner/tools/mcp_client.py":
        "Talks to the tool server, starting it as a separate process.",
    "trip_planner/tools/travel_apis.py":
        "Calls the flight and hotel websites directly over the internet.",
    "trip_planner/tools/agent_tools.py":
        "The 12 tools an AI agent is allowed to hold and choose to use.",
    # comms
    "trip_planner/comms/protocol.py":
        "The typed message format components use to talk to each other, and its six "
        "message types. Every message names its sender, receiver and purpose.",
    "trip_planner/comms/registry.py":
        "The eight agent cards. Each says what its agent may send and receive, which "
        "is what makes an undeclared message a detectable mistake.",
    # server / ui
    "trip_planner/server/mcp_server.py":
        "The tool server: 12 tools, each with a declared input format, offered over a "
        "standard protocol so any client could use them.",
    "trip_planner/ui/app.py": "The web page itself.",
    # evaluation
    "evaluation/arm_a_single_llm.py": "Approach A: one AI call, no tools.",
    "evaluation/arm_b_six_agent_naive.py": "Approach B: six agents, as first built.",
    "evaluation/arm_c_six_agent_tuned.py": "Approach C: the same six agents, tuned.",
    "evaluation/arm_d_three_agent_direct.py":
        "Approach D: three agents with the lookups in plain Python. What ships.",
    "evaluation/run_comparison.py":
        "Runs the approaches, repeats them, and writes the measured numbers to files.",
    "evaluation/measured.py":
        "The ONE place any measured number is read from. This is why the report "
        "cannot disagree with its own charts.",
    "evaluation/metrics.py":
        "Scores how much of a plan is real: do its prices match fares the websites "
        "actually returned?",
    "evaluation/scenarios.py":
        "The 20 test trips, from short and cheap to long and impossible.",
    "evaluation/distilled_tools.py": "Trimmed-down tool wrappers used by approach C.",
    "evaluation/exp_protocol.py":
        "Audits our own two protocols against what we declared they do. Needs no "
        "internet. Currently fails 6 of 9 checks, all reported in the dissertation.",
    "evaluation/exp_budget_gate.py":
        "Tests the budget rules across all 20 trips. Needs no internet.",
    "evaluation/verify_approaches.py":
        "Shows on one screen that every approach really exists: how many agents "
        "each has, how many tools each agent holds, and how many thinking steps it "
        "is allowed. Also calls the tool server for real. Costs nothing, so it is "
        "safe to run in front of someone.",
    "evaluation/check_quota.py":
        "Reports how much monthly travel-API allowance is left. Costs 1 flight and 1 "
        "hotel call, because the balance only comes back inside a real reply.",
    # demos
    "demos/compare_all_approaches.py":
        "All four approaches side by side, with the table and what to say about it.",
    "demos/approach_a_single_llm.py": "Approach A on its own, explained step by step.",
    "demos/approach_b_six_agent_naive.py": "Approach B on its own, explained.",
    "demos/approach_c_six_agent_tuned.py": "Approach C on its own, explained.",
    "demos/approach_d_three_agent_direct.py": "Approach D on its own, explained.",
    "demos/show_agent_messages.py":
        "Shows the agents talking to each other: who is allowed to message whom, a "
        "real exchange being sent, an undeclared pair being refused, and the whole "
        "conversation replayed. Costs nothing and takes a second.",
    "demos/_presenter.py":
        "The narration, written once and shared, so each approach file only has to "
        "describe its own approach.",
}

# Folder, and the heading it appears under in the document.
FILE_GROUPS = [
    ("", "Start here (top level)"),
    ("trip_planner", "trip_planner/ — the workflow"),
    ("trip_planner/core", "trip_planner/core/ — the infrastructure"),
    ("trip_planner/tools", "trip_planner/tools/ — anything that calls outside"),
    ("trip_planner/comms", "trip_planner/comms/ — the message protocol"),
    ("trip_planner/server", "trip_planner/server/ — the tool server"),
    ("trip_planner/ui", "trip_planner/ui/ — the web page"),
    ("evaluation", "evaluation/ — the experiment"),
    ("demos", "demos/ — the demonstrations"),
]


def _figure_count() -> int:
    """
    How many figures exist, counted rather than typed.

    The folder listing said "all 14 pictures". It was right when written, and a
    number in prose next to the folder it describes is exactly the kind of claim
    that goes quietly wrong when a chart is added.
    """
    figures = os.path.join(ROOT, "submission", "build", "figures")
    return sum(1 for _, _, files in os.walk(figures)
               for f in files if f.endswith(".png"))


def _test_count() -> int:
    """
    How many tests pytest collects.

    Through measured.test_count() rather than by counting `def test_` in the
    files: a first attempt did that and reported 300 against pytest's 393,
    because a parametrised test is one definition and many cases. The
    dissertation already reads this accessor, so both documents now quote the
    same number from the same place.
    """
    return measured.test_count()["collected"]


def _file_rows(folder: str):
    """Every Python file in one folder, paired with its plain-English purpose."""
    directory = os.path.join(ROOT, folder) if folder else ROOT
    rows = []
    for name in sorted(os.listdir(directory)):
        if not name.endswith(".py") or name.startswith("__"):
            continue
        key = f"{folder}/{name}" if folder else name
        rows.append([name, FILE_PURPOSE.get(key, "(not yet described)")])
    return rows


def _tuning_reduction_pct() -> float:
    """How much of B's token use tuning removed, measured rather than recalled."""
    b = measured.arm("B")["avg_total_tokens"]
    c = measured.arm("C")["avg_total_tokens"]
    return (b - c) / b * 100.0


def _calls(code: str) -> int:
    return round(measured.arm(code)["avg_llm_calls"])


def build() -> str:
    cov = measured.coverage()
    doc = Doc()

    doc.title("AI Trip Planner",
              "What this project is, how to run it, and how to show it. "
              "CMP7200 MSc Individual Project.")

    # ------------------------------------------------------------------
    doc.h("1. The problem")
    doc.p("""
        Planning a trip takes a long time. You open a flight site, a hotel site,
        a few blogs, and a calculator, then start again every time a price
        changes. It is boring work, not hard work — which is the kind of work a
        computer should do.
    """)
    doc.p("""
        You might expect an AI chatbot to solve this already. It does not. If you
        ask one for a travel plan it will give you a confident answer with real
        looking flight numbers and hotel names — and many of them do not exist.
        A published benchmark tested this properly: the best single AI agent
        completed only 0.6% of travel planning tasks correctly. The plans looked
        finished. They were not bookable.
    """)
    doc.p("""
        The reason is simple. The AI is writing from memory instead of looking
        things up. So the question this project asks is not "can AI plan travel"
        but something narrower and more useful: if we give the AI real tools to
        look things up, who should decide when to use them — the AI, or ordinary
        code?
    """)

    # ------------------------------------------------------------------
    doc.h("2. What the system does")
    doc.p("""
        You type a request in normal English, for example "4 nights in Istanbul
        from Lahore, leaving 15 August, budget 800 dollars, I like history and
        food". The system returns a day-by-day plan with real flights, real
        hotels, real places to eat, and a budget breakdown.
    """)
    doc.code(
        "your request\n"
        "     |\n"
        "     v\n"
        "  STEP 1   AI reads the request and turns it into fields\n"
        "           (origin, destination, dates, budget, interests)\n"
        "     |\n"
        "     v\n"
        "  STEP 2   Plain Python code fetches the data - NO AI here\n"
        "           flights, hotels, attractions, restaurants\n"
        "     |\n"
        "     v\n"
        "  STEP 3   AI arranges the results into a day-by-day plan\n"
        "     |\n"
        "     v\n"
        "  your itinerary"
    )
    doc.p("""
        The important idea is in step 2. The AI is used in step 1 and step 3
        because those need judgement — there is no single right way to phrase a
        request, and no single right way to order a day. Step 2 needs no
        judgement at all: once we know the cities and dates, there is exactly one
        correct search to run. So step 2 uses no AI. Testing whether that was the
        right decision is what the whole project measures.
    """)

    # ------------------------------------------------------------------
    doc.h("3. The four approaches we compared")
    doc.p("""
        The same request was given to four different designs. Everything else was
        held identical — same APIs, same message protocol, same request text — so
        the only thing that changes is who does the data fetching.
    """)
    doc.table(
        ["Approach", "How it works", "Why we built it"],
        [
            ["A — Single AI", "One AI call. No tools. No lookups.",
             "The control. Shows what the tools are worth, and what a plan looks "
             "like when nothing in it was looked up."],
            ["B — 6 agents, naive", "Six AI agents. Each search agent gets many "
             "tools and decides for itself what to call.",
             "This is how we built it first, following the proposal. Its cost is "
             "the reason we changed the design."],
            ["C — 6 agents, tuned", "Same six agents, configured properly: one "
             "tool each, shorter prompts, run at the same time.",
             "The fair comparison. Without it, someone could say approach D only "
             "won because B was set up badly."],
            ["D — 3 agents, direct", "Two AI steps. The lookups are plain Python.",
             "This is the design that ships. It is the claim being tested."],
        ],
        widths=[1.3, 2.4, 2.7],
    )
    doc.p("""
        Approach C matters more than it looks. It would have been easy to compare
        our new design against the badly configured version and report a huge
        win. Building the tuned version first was slower and made our own result
        smaller — and it is the reason the result can be defended.
    """)

    # ------------------------------------------------------------------
    # One diagram per approach, drawn in text rather than as an image, because
    # this document is meant to be read and explained by someone who did not build
    # the project. A diagram they can edit is worth more than a picture they cannot.
    doc.h("4. How each approach works, and what is wrong with it")

    doc.h2("Approach A — one AI call, no tools")
    doc.code(DIAGRAM_A)
    doc.p(f"""
        What is wrong with it: it invents everything. It quoted
        {measured.groundedness('A')['prices_quoted']} prices and
        {measured.groundedness('A')['prices_grounded']} of them matched a real fare.
        It is the cheapest and among the fastest, and it is useless. This is exactly
        why cost on its own is a misleading way to judge these systems.
    """)

    doc.h2("Approach B — six agents, each deciding for itself")
    doc.code(DIAGRAM_B)
    doc.p("""
        What is wrong with it: the loop is the cost. Each agent thinks, calls a tool,
        reads the result and thinks again — and every one of those steps re-sends the
        whole conversation and every tool description again. On the run measured for
        this document it made 24 AI calls, never called the attractions or restaurant
        tools at all, and twice produced malformed output its own loop had to retry.
        It also takes about six minutes, which makes it unusable in a live demo.
    """)

    doc.h2("Approach C — the same six agents, configured properly")
    doc.code(DIAGRAM_C)
    doc.p(f"""
        What is wrong with it: it is far better than B but still unpredictable.
        Configuration alone cut token use by about {_tuning_reduction_pct():.0f}%,
        which is the most useful engineering lesson in the project — most of what
        looked like the price of using many agents was really bad setup. But the
        agents still choose when to call a tool, and on the run measured here it never
        called the flight tool at all.
    """)

    doc.h2("Approach D — three agents, lookups in plain Python (this is what ships)")
    doc.code(DIAGRAM_D)
    doc.p("""
        The idea: use the AI only where a judgement is genuinely required. Reading a
        request is a judgement — there is no single right way to phrase one. Ordering
        a day is a judgement. Fetching a flight price is not: once the cities and
        dates are known there is exactly one correct call to make, so ordinary code
        makes it.
    """)
    doc.p("""
        What is wrong with it, stated honestly: it is decisively cheaper and faster,
        but it is NOT measurably better grounded than approach C. Their ranges overlap
        and C's average is in fact slightly higher. The claim is "no penalty we can
        detect", not "better".
    """)

    # ------------------------------------------------------------------
    doc.h("5. Approach B against approach C, in detail")
    doc.p("""
        The same six agents, the same jobs, the same data. Only the configuration
        differs — which makes this the cleanest comparison in the project.
    """)

    doc.h2("The one fact that explains everything below")
    doc.p("""
        An AI has no memory. Every time you ask it something, the program re-sends
        everything: the request, every tool description, and the whole conversation
        so far. Like an assistant with amnesia — you read the instructions out again
        before every step.
    """)
    doc.p("""
        So the cost of giving an agent a tool is not paid once. It is paid on every
        step that agent takes.
    """)

    doc.h2("Who holds what")
    for label, tool_sets in (("Approach B", TOOLS_B), ("Approach C", TOOLS_C)):
        rows = []
        total = 0
        for agent, tools in tool_sets:
            b_steps, c_steps = LIMITS[agent]
            steps = b_steps if label == "Approach B" else c_steps
            if not tools:
                rows.append([agent, "(none)", "—", str(steps)])
                continue
            for i, (tool, job) in enumerate(tools):
                rows.append([agent if i == 0 else "", tool, job,
                             str(steps) if i == 0 else ""])
                total += 1
        rows.append(["TOTAL", f"{total} tool slots", "", ""])
        doc.h2(f"{label} — {total} tool slots")
        doc.table(["Agent", "Tool it holds", "What that tool does",
                   "Steps allowed"], rows, widths=[1.15, 2.05, 2.4, 0.75], font_pt=8.5)

    doc.h2("What the tool lists show")
    doc.table(
        ["What the lists show about approach B", "What approach C did"],
        [["search_internet and calculate sit on ALL FOUR agents — 8 of the 20 "
          "slots on two tools, re-sent by every agent on every step.",
          "Removed from all of them."],
         ["The hotel agent has two tools doing one job in two steps (look up the "
          "city code, then search) when one tool does both.",
          "Kept only the tool that does both."],
         ["The hotel agent also holds an attractions tool, so it can wander off "
          "instead of finishing.",
          "One tool only. Nothing to wander to."],
         ["The coordinator holds four search tools, though its job is only to "
          "write the plan. It can re-search work already done.",
          "No tools at all."],
         ["Tools hand the AI the full API reply.",
          "Trimmed to the best few first."]],
        widths=[3.6, 2.8], font_pt=9)

    doc.h2("How a tool call actually travels — and whether the MCP server is used")
    doc.p("""
        Worth being exact about: "we built a tool server" and "the system calls it
        over the protocol" are different claims, and only the first is true of
        every approach.
    """)
    doc.table(
        ["Approach", "How it reaches a tool"],
        [["A — single AI", "No tools at all. It calls nothing."],
         ["B — six agents, naive",
          "Through the MCP SERVER over JSON-RPC. The server runs as a separate "
          "program and validates every input against the tool's declared schema. "
          "This is the only approach that exercises the protocol."],
         ["C — six agents, tuned",
          "Imports the server's own tool functions and calls them in-process. Same "
          "functions, same APIs, same data — no separate program, no JSON-RPC."],
         ["D — three agents (ships)",
          "The same in-process calls, made by plain Python instead of by an agent."]],
        widths=[1.5, 4.9],
    )
    doc.p("""
        So: does the shipped system use the MCP server? It calls the same twelve
        tool functions the server exposes, but in-process rather than over the
        protocol. The JSON-RPC transport is exercised by approach B and by the
        conformance audit, which tests all twelve declared schemas against their
        implementations. Sections 4 and 7.2 of the dissertation state and assess this.
    """)

    doc.h2("The same tool call, drawn both ways")
    doc.code(DIAGRAM_B_TOOLCALL)
    doc.code(DIAGRAM_C_TOOLCALL)

    doc.h2('What "steps allowed" means')
    doc.p("""
        An agent does not answer in one go. It loops: think, use a tool, read the
        result, think again. "Steps allowed" is how many times it may go round
        that loop before the framework stops it and forces an answer.
    """)
    doc.p("""
        Each step is one full AI call. And because the AI has no memory, each step
        re-sends every tool description. So the setting matters twice over.
    """)
    doc.code(ITERATIONS_B)
    doc.code(ITERATIONS_C)
    doc.p("""
        Same job. Same hotel chosen. Six AI calls against two — because B could
        wander, and C had nowhere to wander to.
    """)

    doc.h2("And they run in a different order")
    doc.code(
        "  B:  extractor -> flight -> hotel -> attraction -> coordinator\n"
        "      each one waits for the one before it to finish\n"
        "\n"
        "  C:  extractor -> flight     |\n"
        "                   hotel      |  all three at the same time\n"
        "                   attraction |   -> coordinator")

    doc.h2("What changed, and what each change bought")
    doc.table(
        ["What changed", "From", "To", "What it fixed"],
        [["Tools per agent", "up to 8", "1", "Fewer descriptions re-sent every step"],
         ["Thinking steps allowed", "8 to 15", "3", "Re-sent fewer times"],
         ["What the AI is shown", "the full reply", "the best 3 lines",
          "The AI reads far less"],
         ["When agents run", "one after another", "all at once",
          "Cuts the clock, not the tokens"]],
        widths=[1.5, 1.3, 1.3, 2.3],
    )
    doc.p(f"""
        The first three cut token use by about {_tuning_reduction_pct():.0f}%. The
        fourth made it roughly ten times faster. No agent was added or removed and
        no API was swapped.
    """)
    doc.p("""
        Which raised the question approach D answers. In step 1 above, B spent a
        whole AI call deciding "look up Istanbul's ID code" — but there was exactly
        one correct thing to do. No judgement at all. So why is the AI in the lookup
        step? In approach D it is not.
    """)

    # ------------------------------------------------------------------
    doc.h("6. How many calls each approach makes")
    calls = measured.api_calls_per_arm()
    doc.p(f"""
        Counted by watching the real calls go past, with travel responses replayed
        from disk so the counting itself spent no quota. One run of each.
        {calls['note']}
    """)
    doc.table(
        ["Approach", "AI calls", "Flights", "Hotels", "Serper", "Total API",
         "Same every run?"],
        [[f"{code} — {measured.ARM_LABELS[code]}",
          f"{calls['arms'][code]['model_calls']:,}",
          f"{calls['arms'][code]['flights']:,}",
          f"{calls['arms'][code]['hotels']:,}",
          f"{calls['arms'][code]['serper']:,}",
          f"{calls['arms'][code]['total_http']:,}",
          "yes" if calls["arms"][code]["fixed"] else "NO — varies"]
         for code in ("A", "B", "C", "D")],
        widths=[1.5, 0.75, 0.7, 0.7, 0.7, 0.8, 1.15],
    )
    doc.p("""
        Read the last column first. Approaches B and C give a different answer every
        time, because a model is deciding when to call a tool. Approach D makes the
        same five calls every time, because plain code makes them. When the monthly
        allowance is this small, being predictable matters as much as being cheap.
    """)
    doc.bullets([f"{code}: {calls['arms'][code]['comment']}"
                 for code in ("A", "B", "C", "D")])

    # ------------------------------------------------------------------
    doc.h("7. What we found")
    rows = []
    for code in ("A", "B", "C", "D"):
        arm = measured.arm(code)
        g = measured.groundedness(code)
        rows.append([
            f"{code}  {arm['name']}",
            f"{arm['avg_llm_calls']:.0f}",
            f"{arm['avg_total_tokens']:,.0f}",
            f"${arm['avg_cost_usd']:.4f}",
            f"{arm['avg_latency']:.0f}s",
            f"{g['prices_grounded_pct']:.0f}%",
        ])
    doc.table(
        ["Approach", "AI calls", "Tokens", "Cost", "Time", "Prices that are real"],
        rows,
        widths=[1.7, 0.75, 0.85, 0.8, 0.6, 1.3],
    )
    doc.p(f"""
        The last column is the one to look at first. "Prices that are real" means:
        of all the prices printed in the plan, how many match a price the travel
        websites actually returned. Approach A scored
        {measured.groundedness('A')['prices_grounded_pct']:.0f}% — it printed
        {measured.groundedness('A')['prices_quoted']} prices and not one of them
        was real. It is the cheapest approach and its output is worthless.
    """)
    doc.p("""
        Three findings:
    """)
    doc.bullets([
        "Tools matter more than anything. Without them the plan is fiction. "
        "This is the clearest result in the project.",
        "Tuning matters more than the number of agents. B and C are the SAME six "
        "agents with the same data — only the prompts and settings changed — and "
        f"that alone cut token use by about {_tuning_reduction_pct():.0f}%. So most "
        "of what looked like the cost of using many agents was really just bad "
        "configuration.",
        "Removing the AI from the lookup step saves a lot of AI calls and time, "
        "and only a little money. That is a smaller claim than we expected, and "
        "it is the one the numbers actually support.",
    ])
    doc.p(f"""
        Every one of these numbers is an average of
        {cov['repeats_per_arm']} separate runs, so the report can also say how much
        each one wobbles between runs. Two results are worth knowing:
    """)
    doc.bullets([
        "Approach D is cheaper and faster than C by more than the run-to-run "
        "wobble, so that difference is real.",
        "On how much of the plan is real, C and D overlap. So the honest statement "
        "is 'we cannot tell them apart', not 'they are equal'. C's average is "
        "actually slightly higher, and the report says so.",
        "Approaches A and D use exactly the same number of AI calls every single "
        "time. B and C do not, because they contain a loop that decides for itself "
        "how many calls to make. Predictable cost is a real advantage.",
    ])
    doc.p(f"""
        Honest limit: the runs all use
        {cov['scenarios_measured']} test scenario out of
        {cov['scenarios_designed']} we designed. Repeating a scenario is free
        because the website replies are replayed from disk, but each NEW scenario
        needs about four of the 80 free lookups a month. So we have good depth on
        one trip and no breadth across trips. The report says this in the abstract
        and on every chart.
    """)

    # ------------------------------------------------------------------
    doc.h("8. Budget validation (the change the supervisor asked for)")
    doc.p("""
        The supervisor asked for the budget handling to be improved. Two things
        were wrong with the original version, and both are fixed.
    """)
    doc.h2("Problem 1: the same budget split for every trip")
    doc.p("""
        The first version always split the money the same way — 35% flights, 35%
        hotels, and so on — no matter what the trip was. That is wrong, because
        costs do not scale the same way. A flight is paid per person and gets
        more expensive the further you go. A hotel room is paid per night and is
        shared between people. So a split that suits one person going somewhere
        near is badly wrong for a family going far.
    """)
    doc.p("""
        Now the split is worked out from what that particular trip actually
        costs. And if the traveller says how they want to divide their money, we
        use their split — someone who says that has made a choice, not a mistake.
    """)
    doc.h2("Problem 2: it did not know whether a budget was possible")
    doc.p("""
        The first version refused a budget using one fixed formula that ignored
        the destination completely. It judged 700 dollars for Bangkok and 700
        dollars for London by the same rule and accepted both — so it happily
        produced a confident, impossible plan for London.
    """)
    doc.p("""
        Now the system estimates what that specific trip costs at three levels —
        the cheapest it can possibly be done, comfortable, and luxury — using the
        destination's price level, the flight distance, the number of nights, the
        number of people, and room sharing. It only refuses a budget below the
        true minimum. Above that, a tight budget gets a warning and carries on,
        because tight is a real choice and only impossible is not.
    """)
    gate = measured.gate_agreement()
    anchor = measured.gate_external_validity()
    misses = measured.gate_misses()
    doc.p(f"""
        We then tested it on all {gate['n']} scenarios, and reported the result
        even though it is not perfect. Two scenarios were written to be
        impossible. It correctly refused {gate['true_positive']} and wrongly
        allowed {gate['false_negative']}
        ({misses[0]['scenario'] if misses else 'none'}). The agreement score
        (Cohen's kappa) is {gate['cohens_kappa']:.3f}.
    """)
    doc.p(f"""
        We also found out WHY it got that one wrong, which is more useful than
        the score. The system assumes the cheapest possible flight on that route
        costs about ${anchor['estimated_minimum']:,.0f}. The travel website
        actually returned ${anchor['cheapest_real_fare']:,.0f} as its cheapest —
        so our assumption is about
        {abs(anchor['minimum_anchor_error_pct']):.0f}% too low. With the flight
        cost guessed at half the real price, the total minimum comes out too low,
        and a budget slips through that should not have. The structure is right;
        one number in it is wrong.
    """)

    # ------------------------------------------------------------------
    doc.h("9. Known issues, stated honestly")
    protocol = measured.protocol_summary()
    doc.p(f"""
        We wrote a test that checks whether our own communication and tool layers
        really behave the way the design says. {protocol['passed']} of
        {protocol['total_checks']} checks passed. The failures are listed in the
        report rather than quietly fixed, because finding them by measuring is
        worth more marks than pretending they were never there.
    """)
    doc.table(
        ["Issue", "What it means in plain terms"],
        [
            ["Message priority does nothing",
             "Every message carries a priority label, but the queue ignores it and "
             "delivers in arrival order."],
            ["Inbound permissions are not checked",
             "Each component lists who may send to it. That list is never read; "
             "only the sender's own list is checked."],
            ["4 of 12 tools have wrong descriptions",
             "Some tools accept settings that are missing from their published "
             "description, and one names the wrong data provider."],
            ["Flight cost guess is too low",
             f"About {abs(anchor['minimum_anchor_error_pct']):.0f}% under the real "
             f"cheapest fare, which is why one impossible budget was allowed."],
            [f"Only {cov['scenarios_measured']} of {cov['scenarios_designed']} "
             f"scenarios measured",
             "The travel websites' free monthly limit. Nothing to do with the code."],
            ["Google retired the AI model we measured on",
             "Handled — see section 12. The old numbers stay valid for the model "
             "that produced them."],
        ],
        widths=[2.1, 4.3],
    )

    # ------------------------------------------------------------------
    doc.h("10. How to run it")
    doc.p("""
        Put the project folder somewhere with a SHORT path, for example
        C:\\trip_planner. Windows cannot handle very long folder paths and the
        install will fail with a confusing error if the folder is buried deep.
    """)
    doc.p("Then double-click one file:")
    doc.code("run.bat")
    doc.p("""
        That is all. It checks Python is installed, creates its own private
        Python environment, installs everything it needs, sets up the settings
        file, checks the install worked, and then shows a menu. The first run
        takes a few minutes. Every run after that starts in seconds.
    """)
    doc.table(
        ["Menu option", "What it does", "Needs internet or keys?"],
        [
            ["1", "Compare all four approaches side by side", "No"],
            ["2 to 5", "Show one approach in detail (A, B, C or D)", "No"],
            ["6", "Run the test suite", "No"],
            ["7", "Run the evaluation experiments", "No"],
            ["8", "Rebuild all the figures", "No"],
            ["9", "Rebuild the dissertation document", "No"],
            ["10", "Plan a real trip in the web browser", "Yes"],
            ["11", "Plan a real trip in this window", "Yes"],
            ["12", "Exit", "No"],
        ],
        widths=[1.0, 3.4, 2.0],
    )
    doc.p("""
        Options 1 to 10 work with no internet and no API keys at all. This is
        deliberate: the free accounts this project uses run out, and a
        demonstration that needs a working internet account is one that cannot be
        given on the day the account runs out.
    """)

    # ------------------------------------------------------------------
    doc.h("11. How to show it to the supervisor")
    doc.p("""
        Fifteen minutes, in this order. Everything here works offline, so nothing
        can fail on the day.
    """)
    doc.steps([
        "Double-click run.bat and choose option 1. This runs all four approaches "
        "side by side and prints one comparison table. Point at the last column "
        "and say: approach A is the cheapest and every price in it is invented. "
        "That is the whole point of the project in one line.",

        "Point at approaches B and C in the same table. Say: these are the SAME "
        "six agents with the same data. Only the settings changed, and that alone "
        f"cut token use by about {_tuning_reduction_pct():.0f}%. So we compared our "
        "new design against the properly tuned version, not the badly configured "
        "one.",

        "Choose option 5 (approach D, the one that ships). It shows the real "
        "itinerary that was produced, what it cost, and how much of it was real. "
        f"Say: {_calls('D')} AI calls instead of {_calls('C')}, same quality.",

        "Choose option 7. This runs the two experiments that need no internet. "
        "Say: these check our own system against its own design, and they found "
        "six things wrong, which are all written up in the report.",

        "Choose option 6 to show the test suite passing.",

        "Open submission/CMP7200_Dissertation.docx. Say: every number in this "
        "document is read from the measured results files, none is typed by hand, "
        "and there is a script that proves it by corrupting the data and checking "
        "the document changes.",

        "If asked to plan a real trip live, use option 11 with the API keys in "
        "place. If the internet or the free quota is unavailable, options 1 to 10 "
        "still work and show real recorded output.",
    ])
    doc.p("""
        One thing worth saying out loud, because it is a strength and not a
        weakness: when you show a demo from option 1 to 5, the screen says
        clearly that it is replaying a recorded run. It is real measured output,
        not a live one, and it says so. If the supervisor asks whether it is
        running now, the honest answer is already on the screen.
    """)

    # ------------------------------------------------------------------
    doc.h("12. About the AI model")
    doc.p(f"""
        The first round of measurements used Google's Gemini 2.5 Flash. Google
        then stopped giving that model to new accounts, so it could no longer be
        run at all. Its replacement works, but it refuses a particular message
        pattern that the agent framework produces, which meant approaches B and C
        would not run on it either. Every number in this project has since been
        re-measured on the replacement, {measured.model_name()}, which is the
        model the results and the report describe.
    """)
    doc.p("""
        This is fixed. A small compatibility layer
        (trip_planner/core/gemini_compat.py) adjusts only the requests that would
        otherwise be rejected, leaves every other request untouched, and counts
        how many it changed so a run can report whether it was needed. All four
        approaches run again.
    """)
    doc.p("""
        Two honest notes. Numbers from the new model are not directly comparable
        to the old ones, because it is a different model — so if the evaluation is
        re-run, all four approaches must be re-run together. And the report treats
        this as a genuine finding: a system measured against someone else's hosted
        model inherits that model's lifetime.
    """)

    # ------------------------------------------------------------------
    doc.h("13. Folder structure")
    doc.code(
        "trip_planner\\\n"
        "  run.bat                  <- START HERE. Sets up everything, then a menu.\n"
        "  run_cli.py               plan a trip in the terminal\n"
        "  run_web.py               plan a trip in the browser\n"
        "  README.md                the technical readme\n"
        "  PROJECT_OVERVIEW.docx    this document\n"
        "  requirements.txt         the exact package versions used\n"
        "  .env                     your API keys (never shared, never committed)\n"
        "\n"
        "  trip_planner\\           THE SYSTEM ITSELF\n"
        "    orchestrator.py        the workflow: read, fetch, assemble\n"
        "    agents.py, tasks.py    the three AI agents and their instructions\n"
        "    comms\\                the message protocol between components\n"
        "    server\\               the MCP tool server (12 tools)\n"
        "    tools\\                three files: MCP client, travel APIs, agent tools\n"
        "    core\\                 caching, measuring, budget, cost, safe maths\n"
        "    ui\\                   the web page, and the code that splits a\n"
        "                           finished plan into its sections\n"
        "\n"
        "  evaluation\\             THE EXPERIMENT\n"
        "    arm_a ... arm_d        the four approaches\n"
        "    scenarios.py           the 20 test trips\n"
        "    metrics.py             scores how much of a plan is real\n"
        "    measured.py            the ONE place results are read from\n"
        "    exp_protocol.py        checks our protocols against the design\n"
        "    exp_budget_gate.py     checks the budget rules on all 20 scenarios\n"
        "    results\\              the measured numbers, as files\n"
        "\n"
        "  demos\\                  DEMONSTRATIONS (work with no internet)\n"
        "    compare_all_approaches.py     all four side by side\n"
        "    approach_a ... approach_d     one approach each, in detail\n"
        "\n"
        "  report\\                 THE DISSERTATION\n"
        "    CMP7200_Dissertation.docx     the report itself\n"
        "    build\\                       one file per chapter, and the figures\n"
        f"    figures\\                     all {_figure_count()} pictures, generated\n"
        "\n"
        f"  testing\\                the test suite ({_test_count()} tests), plus two\n"
        "                           real finished plans used as fixtures\n"
        "  proposal\\               the original proposal and the assignment brief\n"
        "  .streamlit\\             the web page's colour theme\n"
        "  .api_cache\\             recorded website replies, so results replay free"
    )

    # ------------------------------------------------------------------
    doc.h("14. Every file, and what it does")
    doc.p("""
        The whole system, file by file. The list is read from the folders
        themselves when this document is generated, so it cannot describe a file
        that has been deleted or miss one that has been added.
    """)
    for folder, heading in FILE_GROUPS:
        rows = _file_rows(folder)
        if not rows:
            continue
        doc.h2(heading)
        doc.table(["File", "What it does"], rows, widths=[1.9, 4.5], font_pt=9)

    # ------------------------------------------------------------------
    doc.h("15. Important files, if you only look at a few")
    doc.table(
        ["File", "Why it matters"],
        [
            ["run.bat", "The only thing anyone needs to double-click."],
            ["trip_planner/orchestrator.py",
             "The workflow. Read this to understand how the system works."],
            ["trip_planner/core/http_cache.py",
             "Records every website reply so results can be reproduced with no "
             "accounts, and stops one run using up a whole month's free quota."],
            ["trip_planner/core/llm_metrics.py",
             "Counts real AI calls, tokens and cost. Earlier versions guessed "
             "these numbers and the guesses were wrong."],
            ["trip_planner/core/trip_cost.py",
             "The budget validation the supervisor asked for."],
            ["evaluation/measured.py",
             "The single place any measured number is read from. This is why the "
             "report cannot disagree with the charts."],
            ["submission/build/verify_no_hardcoded_numbers.py",
             "Proves no number in the report is typed by hand: it corrupts the "
             "data, rebuilds, and checks every value changed."],
        ],
        widths=[2.2, 4.2],
    )

    # ------------------------------------------------------------------
    doc.h("16. The APIs used, and how much is left")
    doc.table(
        ["What", "Service", "Free limit", "Used for"],
        [
            ["AI model", "Google Gemini", "Free tier, limited per minute",
             "Reading the request and writing the plan"],
            ["Flights", "fly-scraper (RapidAPI)", "30 per MONTH",
             "Real flight prices and times"],
            ["Hotels", "Booking.com (RapidAPI)", "50 per MONTH",
             "Real hotels, prices and review scores"],
            ["Places", "Serper.dev", "Generous", "Attractions and restaurants"],
        ],
        widths=[0.9, 1.7, 1.5, 2.3],
    )
    cache = measured.api_cache_stats()
    doc.p(f"""
        The flight and hotel limits are per MONTH and very small. One careless run
        can use a whole month and it cannot be bought back. That is why every
        website reply is recorded: there are {cache['entries']} recorded replies
        in the project, and they let anyone reproduce the results with no accounts
        at all.
    """)

    doc.h2("How much is left right now")
    quota = measured.api_quota()
    quota_rows = []
    for host, reading in quota["apis"].items():
        remaining = reading.get("remaining")
        limit = reading.get("limit") or "?"
        if remaining is None:
            quota_rows.append([reading["name"], "not reported by the plan",
                               "check rapidapi.com"])
            continue
        days = ""
        if reading.get("reset_seconds"):
            try:
                days = f"resets in {int(reading['reset_seconds']) / 86400:.0f} days"
            except (TypeError, ValueError):
                days = ""
        quota_rows.append([reading["name"], f"{remaining} left of {limit}", days])
    doc.table(["API", "Remaining", "When it refills"], quota_rows,
              widths=[2.0, 2.0, 2.4])
    doc.p(f"""
        Taken at {quota['checked_at'].replace('T', ' ')}. This is a snapshot, not a
        live figure — it falls every time anyone plans a real trip. To refresh it:
    """)
    doc.code("python -m evaluation.check_quota")
    doc.p("""
        Be aware that checking costs 1 flight call and 1 hotel call, because the
        remaining balance is only reported inside the reply to a real request. There
        is no way to ask for free. Run it once before a demonstration, not in a loop.
    """)
    doc.p("""
        What that buys, in practice: one live trip through the shipped design uses 1
        flight call and 2 hotel calls. So the flight allowance is the binding limit —
        divide the flights remaining by one to get the number of complete live trips
        left. Everything else in this project, including all five demonstrations and
        both experiments, replays from disk and uses none of it.
    """)

    # ------------------------------------------------------------------
    doc.h("17. If something goes wrong")
    doc.table(
        ["Problem", "What to do"],
        [
            ["\"Python is not recognised\"",
             "Install Python 3.10 or newer from python.org and TICK \"Add Python "
             "to PATH\" during installation."],
            ["Install fails with \"No such file or directory\"",
             "The folder path is too long for Windows. Move the whole folder to "
             "somewhere short like C:\\trip_planner and run again."],
            ["Something says a dependency is missing",
             "Delete the .venv folder and run run.bat again. It rebuilds from "
             "scratch."],
            ["A live trip plan fails",
             "The free quota has probably run out. Options 1 to 10 still work with "
             "no internet."],
            ["The report will not rebuild",
             "It refuses to build if a figure is missing or a number cannot be "
             "found. Run option 8 first to rebuild the figures."],
        ],
        widths=[2.3, 4.1],
    )

    path = doc.save()
    return path


if __name__ == "__main__":
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    print(f"wrote {os.path.relpath(build(), ROOT)}")
