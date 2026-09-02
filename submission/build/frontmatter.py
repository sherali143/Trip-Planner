"""Title page, abstract, contents, and the lists of figures and tables."""

from __future__ import annotations

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from trip_planner.evaluation import measured
from submission.build.common import (ACCENT, AWARD, FACULTY, MODULE_CODE,
                                MODULE_TITLE, MUTED, STUDENT_NUMBER,
                                SUBMISSION_DATE, TITLE, UNIVERSITY, Report, val)


def _centred(report: Report, text: str, *, size: float = 11, bold: bool = False,
             colour=None, space_after: float = 6):
    paragraph = report.doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if colour is not None:
        run.font.color.rgb = colour
    report._count(text)
    return run


def title_page(report: Report) -> None:
    report.start_excluded("title page")
    _centred(report, UNIVERSITY, size=13, bold=True, space_after=2)
    _centred(report, FACULTY, size=10.5, colour=MUTED, space_after=28)

    _centred(report, TITLE, size=17, bold=True, colour=ACCENT, space_after=26)

    _centred(report, f"{MODULE_CODE}  {MODULE_TITLE}", size=11.5, space_after=2)
    _centred(report, "Assessment 2: Project Dissertation", size=11.5, space_after=26)

    # Student number only. The brief: "This assessment will be marked
    # anonymously and should show your student number only."
    _centred(report, f"Student Number: {STUDENT_NUMBER}", size=12.5, bold=True,
             space_after=26)

    _centred(report,
             f"Submitted in partial fulfilment of the requirements for the degree of "
             f"{AWARD}", size=10.5, colour=MUTED, space_after=6)
    _centred(report, SUBMISSION_DATE, size=10.5, colour=MUTED, space_after=18)

    # Filled in by declare_word_count() once the body has been built, because
    # the count is not known while this page is being written. Without it a
    # marker running Word's own count sees the whole file — front matter,
    # references and every appendix — against a limit that applies to the eight
    # chapters only.
    report.wordcount_run = _centred(report, "Word count: pending",
                                    size=10.5, colour=MUTED, space_after=30)

    coverage = measured.coverage()
    _centred(report,
             f"All quantitative results in this dissertation were produced by "
             f"{coverage['model']} across "
             f"{val(coverage['scenarios_measured'])} of "
             f"{val(coverage['scenarios_designed'])} designed evaluation scenarios, "
             f"with the API layer in {coverage['api_mode']} mode. Every figure and "
             f"table is generated from the committed result files; no value in this "
             f"document is typed by hand.",
             size=9, colour=MUTED)


def abstract(report: Report) -> None:
    report.start_excluded("abstract")
    report.unnumbered_h1("Abstract")

    coverage = measured.coverage()
    d_vs_c = measured.improvement("D_vs_C")
    c_vs_b = measured.improvement("C_vs_B")
    control = measured.groundedness("A")
    tuned = measured.groundedness("C")
    direct = measured.groundedness("D")
    protocol = measured.protocol_summary()
    gate = measured.gate_agreement()
    anchor = measured.gate_external_validity()

    report.p(f"""
        Travel planning tests whether a language model can be trusted with a task
        that has verifiable answers: a plan names flights, hotels, prices and
        dates, and each either corresponds to something bookable or does not.
        Single-model assistants perform badly on exactly this — the strongest
        single-agent baseline on the TravelPlanner benchmark completes 0.6% of
        tasks, failing largely through invented venues and fabricated prices
        [@xie2024]. This dissertation asks a narrower question: given a
        schema-validated tool layer and a typed inter-agent protocol, what does
        delegating retrieval to a model cost, and what does it buy?
    """)

    report.p(f"""
        A working system was built and evaluated against itself in four
        configurations: one model with no tools; a six-agent design as first
        implemented; the same six agents after their prompt economics were tuned;
        and a three-agent design keeping both protocols but retrieving in ordinary
        Python. All four share one Model Context Protocol server exposing twelve
        tools over three live travel APIs, and one typed message layer. Every
        model request is counted by provider callback rather than estimated, and
        every HTTP response recorded, so results replay without an API key.
    """)

    report.p(f"""
        Three findings matter, and the second is not the expected one. Retrieval
        delegated to a model costs more without being better: against the tuned
        six-agent arm, the three-agent design used
        {val(d_vs_c['llm_calls_pct'], '{:.1f}')}% fewer requests,
        {val(d_vs_c['latency_pct'], '{:.1f}')}% less time and
        {val(d_vs_c['cost_pct'], '{:.1f}')}% less money, with cost and latency
        intervals that do not overlap over {val(coverage['repeats_per_arm'])} runs
        of each. Its groundedness interval does overlap, so the saving carries no
        quality penalty this evidence can detect — though the tuned arm's mean is
        the higher, recorded rather than rounded away. Second, most of the
        multi-agent penalty was implementation rather than architecture: tuning
        alone cut token use by {val(c_vs_b['tokens_pct'], '{:.1f}')}%, which
        narrows the defensible claim considerably. Third, tool access separates a
        plausible itinerary from a usable one — the tool-less arm quoted
        {val(control['prices_quoted'])} prices and matched
        {val(control['prices_grounded'])} to anything real.
    """)

    report.p(f"""
        The evaluation also turned on the artefact. A conformance audit of its own
        protocol layer passed {val(protocol['passed'])} of
        {val(protocol['total_checks'])} checks: message priority is declared on
        every message and never honoured, and
        {val(len(measured.mcp_schema_stats()['defective_tools']))} of
        {val(measured.mcp_schema_stats()['tools_total'])} tool schemas disagree
        with their implementations. The budget feasibility gate reached a Cohen's
        kappa of {val(gate['cohens_kappa'], '{:.3f}')} across all twenty designed
        scenarios, and the reason is traceable: its cheapest-fare anchor sits
        {val(abs(anchor['minimum_anchor_error_pct']), '{:.0f}')}% below the
        cheapest fare the API returned for the one route with recorded fares.
        These defects were found by measuring the artefact, and they are the most
        useful evidence it produced.
    """)

    report.p(f"""
        The principal limitation is breadth rather than depth. Free-tier quotas on
        the flight and hotel APIs mean the four-arm comparison rests on
        {val(coverage['scenarios_measured'])} recorded scenario of
        {val(coverage['scenarios_designed'])} designed. Repeats replay recorded
        responses and cost no quota, which is why every figure carries an interval,
        but additional scenarios cannot be bought the same way: differences on this
        scenario are measurable, and generalisation to other trips is not
        established. The protocol and budget-gate experiments are quota-free and
        cover all twenty. The contribution is a reproducible harness and a narrow
        measured finding, not a benchmark result.
    """)


def contents(report: Report) -> None:
    """
    Insert a table-of-contents field for Word to populate.

    python-docx cannot compute page numbers, so the field is written as raw
    WordprocessingML and Word fills it in on first open. This is flagged in the
    build output because it needs one manual action.
    """
    report.start_excluded("contents")
    report.unnumbered_h1("Contents")

    paragraph = report.doc.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = ("Right-click here and choose 'Update Field' to build the "
                        "table of contents.")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, placeholder, end):
        run._r.append(element)


def appendix_title_page(report: Report) -> None:
    """
    Cover page for the appendices when they are submitted as a separate file.

    Carries the student number and nothing else identifying, for the same
    anonymous-marking reason as the dissertation's own title page, and says
    plainly which document it belongs to so the pair cannot be separated by
    accident.
    """
    report.start_excluded("appendix title page")
    _centred(report, UNIVERSITY, size=13, bold=True, space_after=2)
    _centred(report, FACULTY, size=10.5, colour=MUTED, space_after=28)
    _centred(report, "Appendices", size=17, bold=True, colour=ACCENT,
             space_after=10)
    _centred(report, TITLE, size=11.5, colour=MUTED, space_after=26)
    _centred(report, f"{MODULE_CODE}  {MODULE_TITLE}", size=11.5, space_after=2)
    _centred(report, "Assessment 2: Project Dissertation", size=11.5,
             space_after=26)
    _centred(report, f"Student Number: {STUDENT_NUMBER}", size=12.5, bold=True,
             space_after=26)
    _centred(report,
             "This document accompanies CMP7200_Dissertation.docx. The appendices "
             "are excluded from the word count and are referenced by letter from "
             "the chapters of that document.",
             size=9.5, colour=MUTED)


def _short(caption: str) -> str:
    """
    The identifying part of a caption, for the lists of figures and tables.

    These lists exist so a reader can find a figure, not read about it — the
    full caption is under the figure itself, and repeating all of it here said
    everything twice. Keeps the first sentence, and trims that at a clause
    boundary if it is still long.
    """
    head = caption.split(". ")[0].rstrip(".")
    if len(head.split()) <= 11:
        return head
    for mark in ("; ", ", "):
        if mark in head:
            candidate = head.split(mark)[0]
            if 3 <= len(candidate.split()) <= 11:
                return candidate
    return " ".join(head.split()[:11])


def figure_and_table_lists(report: Report) -> None:
    """
    Lists of figures and tables, generated from what was actually inserted.

    Written after the body is built, so the lists cannot disagree with the
    document. Called last by build_report.py and moved into position by Word's
    own field update is not possible, so these appear at the end of the front
    matter section instead — a compromise noted in the build output.
    """
    report.start_excluded("lists of figures and tables")
    for heading, index in (("Figures", report.figure_index),
                           ("Tables", report.table_index)):
        report.unnumbered_h1(heading, page_break=(heading == "Figures"))
        for tag, caption in index:
            line = f"{tag}   {_short(caption)}"
            paragraph = report.doc.add_paragraph(line)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(3)
            # Counted, even though it is excluded from the limit. These 35 lines
            # are the longest text in the document that no counter saw, and a
            # build that reports its own totals should not be missing a third of
            # the excluded words.
            report._count(line)


def declare_word_count(report: Report) -> None:
    """
    Write the main-body word count onto the title page.

    Called after the body is assembled, so the figure is the counter's own and
    cannot drift from the document it describes. Says what is excluded as well
    as the number, because the number alone invites the wrong comparison.
    """
    if report.wordcount_run is None:      # title page not built
        return

    def sentence(excluded: int) -> str:
        # Says what the figure includes as well as what it leaves out. An
        # earlier wording claimed tables and captions were excluded when the
        # counter counts both — understating what the number covered.
        return (f"Word count: {report.body_words:,} words (Chapters 1–8, "
                f"including headings, lists, tables and captions). Excluded: "
                f"title page, abstract, contents, references, appendices and "
                f"the lists of figures and tables — {excluded:,} words.")

    # This line replaces a placeholder that was itself counted, and it reports
    # the excluded total, so the total depends on the line's own length. One
    # pass settles it: the sentence has the same number of words whatever the
    # digits are, so the length is known before the final figure is.
    placeholder = len(report.wordcount_run.text.split())
    report.excluded_words += len(sentence(0).split()) - placeholder
    report.wordcount_run.text = sentence(report.excluded_words)
